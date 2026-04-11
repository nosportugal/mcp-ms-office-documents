"""Dynamic registration of DOCX template MCP tools.

Similar to dynamic email templates, this module allows defining custom DOCX templates
with placeholders ({{placeholder}}) and YAML configuration for template-specific arguments.

Placeholders in DOCX templates use Mustache syntax:
  - {{placeholder}} - replaced with markdown-formatted text
  - Text supports inline markdown: **bold**, *italic*, `code`, [links](url)

YAML configuration example:
```yaml
templates:
  - name: formal_letter
    description: Formal business letter template
    docx_path: letter_template.docx  # filename only, searched in custom/default templates
    annotations:
      title: Formal Letter Generator
    args:
      - name: recipient_name
        type: string
        description: Full name of the recipient
        required: true
      - name: body
        type: string
        description: Main body text (supports markdown formatting)
        required: true
```
"""
from __future__ import annotations

import io
import re
import logging
from pathlib import Path
from typing import Any, Dict, Optional, Literal

import yaml
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table
from pydantic import Field, create_model
from fastmcp import FastMCP

from upload_tools import upload_file
from template_utils import find_file_in_template_dirs
from .helpers import (
    parse_inline_formatting,
    contains_block_markdown,
    process_markdown_block,
)
from fastmcp.exceptions import ToolError


__all__ = ["register_docx_template_tools_from_yaml"]

logger = logging.getLogger(__name__)

# Type mapping for YAML -> Python types
TYPE_MAP = {
    "string": str, "str": str,
    "int": int, "integer": int,
    "float": float,
    "bool": bool, "boolean": bool,
    "list": list[str], "list[str]": list[str], "list[string]": list[str],
}

# Regex to find Mustache-style placeholders: {{name}} or {{{name}}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{\{?([a-zA-Z_][a-zA-Z0-9_]*)\}?\}\}')



def _insert_markdown_content_after_paragraph(
    doc: DocxDocument,
    paragraph: Paragraph,
    content: str
) -> None:
    """Insert markdown content (including lists and headings) after a paragraph.

    This function processes block-level markdown content including:
    - Headings (# heading)
    - Regular paragraphs with inline formatting
    - Ordered lists (1. item)
    - Unordered lists (- item, * item, + item)

    Args:
        doc: The Word document
        paragraph: The paragraph after which to insert content
        content: The markdown content to insert
    """
    try:
        lines = content.split('\n')
        i = 0

        # Find the paragraph's position in the document body
        body = doc._body._body
        p_element = paragraph._p
        para_idx = list(body).index(p_element)

        # Track how many elements we've inserted
        inserted_count = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Use shared helper to process the markdown block
            i, new_elements = process_markdown_block(doc, lines, i, return_element=True)

            # Insert elements at the correct position
            for elem in new_elements:
                body.insert(para_idx + 1 + inserted_count, elem)
                inserted_count += 1
    except Exception as e:
        logger.error("Failed to insert markdown content after paragraph: %s", e, exc_info=True)


def find_docx_template_by_name(filename: str) -> Optional[str]:
    """Find a specific DOCX template by filename in custom/default template directories.

    Args:
        filename: The filename of the DOCX template (e.g., 'letter_template.docx')

    Returns:
        Absolute path to the template file as string, or None if not found.
    """
    found = find_file_in_template_dirs(filename)
    return str(found) if found else None


def _replace_placeholder_in_paragraph(
    paragraph: Paragraph,
    placeholder: str,
    value: str,
    doc: DocxDocument = None
) -> bool:
    """Replace a placeholder in a paragraph with markdown-formatted text.

    This function handles the case where a placeholder might be split across multiple runs
    (which Word often does when editing documents).

    For block-level content (lists), the content is inserted as new paragraphs after
    the current paragraph.

    Args:
        paragraph: The paragraph to search and modify
        placeholder: The placeholder text including braces (e.g., '{{name}}')
        value: The replacement value (supports markdown formatting)
        doc: The Word document (required for block-level content like lists)

    Returns:
        True if replacement was made, False otherwise
    """
    try:
        # First, try to find the placeholder in the full paragraph text
        full_text = paragraph.text
        if placeholder not in full_text:
            return False

        # Collect all runs and their text
        runs = list(paragraph.runs)
        if not runs:
            return False

        # Build a map of character positions to runs
        combined_text = ""
        run_info = []  # List of (start_pos, end_pos, run)

        for run in runs:
            start = len(combined_text)
            combined_text += run.text
            end = len(combined_text)
            run_info.append((start, end, run))

        # Find the placeholder in the combined text
        placeholder_start = combined_text.find(placeholder)
        if placeholder_start == -1:
            return False

        placeholder_end = placeholder_start + len(placeholder)

        # Store formatting from the run where placeholder starts
        formatting_run = None
        for start, end, run in run_info:
            if start <= placeholder_start < end:
                formatting_run = run
                break

        font_name = formatting_run.font.name if formatting_run else None
        font_size = formatting_run.font.size if formatting_run else None
        font_color_rgb = formatting_run.font.color.rgb if formatting_run else None
        font_color_theme = formatting_run.font.color.theme_color if formatting_run else None

        # Strategy: Rebuild the paragraph content
        # 1. Get text before placeholder
        # 2. Get replacement content (parsed markdown)
        # 3. Get text after placeholder

        text_before = combined_text[:placeholder_start]
        text_after = combined_text[placeholder_end:]

        # Check if the value contains block-level content (lists, headings)
        has_block_content = contains_block_markdown(value)

        # Clear all existing runs
        p_element = paragraph._p
        for run in runs:
            p_element.remove(run._r)

        # Add text before placeholder (plain text, preserve any formatting would be complex)
        if text_before:
            paragraph.add_run(text_before)

        if has_block_content and doc is not None:
            # Insert block content after this paragraph
            _insert_markdown_content_after_paragraph(doc, paragraph, value)

            # If there's text after, add it as a run to this paragraph
            if text_after:
                paragraph.add_run(text_after)

            # If the placeholder occupied the whole paragraph (no surrounding text),
            # the paragraph is now empty – remove it to avoid a spurious blank line.
            if not text_before and not text_after:
                p_element.getparent().remove(p_element)
        else:
            # Simple inline replacement
            # Parse and add the replacement value with markdown formatting
            parse_inline_formatting(value, paragraph)

            # Apply font formatting to newly added runs (from replacement)
            if font_name or font_size or font_color_rgb or font_color_theme:
                # Get runs added after text_before
                new_runs = list(paragraph.runs)
                start_idx = 1 if text_before else 0
                for run in new_runs[start_idx:]:
                    if font_name and not run.font.name:
                        run.font.name = font_name
                    if font_size and not run.font.size:
                        run.font.size = font_size
                    if font_color_rgb and not run.font.color.rgb:
                        run.font.color.rgb = font_color_rgb
                    elif font_color_theme and not run.font.color.theme_color:
                        run.font.color.theme_color = font_color_theme

            # Add text after placeholder
            if text_after:
                paragraph.add_run(text_after)

        return True

    except Exception as e:
        logger.error("Failed to replace placeholder '%s' in paragraph: %s", placeholder, e, exc_info=True)
        return False


def _replace_placeholders_in_paragraph(
    paragraph: Paragraph,
    context: Dict[str, str],
    doc: DocxDocument = None
) -> None:
    """Replace all placeholders in a paragraph with their values.

    This function iteratively replaces placeholders one at a time, re-scanning
    the paragraph after each replacement to handle position shifts correctly.

    Args:
        paragraph: The paragraph to process
        context: Dictionary mapping placeholder names to their values
        doc: The Word document (required for block-level content like lists)
    """
    # Keep replacing until no more placeholders are found
    max_iterations = 100  # Safety limit to prevent infinite loops
    iteration = 0

    while iteration < max_iterations:
        iteration += 1

        # Get current paragraph text and find placeholders
        full_text = paragraph.text
        matches = PLACEHOLDER_PATTERN.findall(full_text)

        if not matches:
            break

        # Find the first placeholder that exists in context
        replaced = False
        for placeholder_name in matches:
            if placeholder_name not in context:
                continue

            value = context[placeholder_name]
            if value is None:
                value = ""
            else:
                value = str(value)

            # Try triple brace first, then double brace
            for placeholder in [f'{{{{{{{placeholder_name}}}}}}}', f'{{{{{placeholder_name}}}}}']:
                if placeholder in paragraph.text:
                    if _replace_placeholder_in_paragraph(paragraph, placeholder, value, doc):
                        replaced = True
                        break

            if replaced:
                break  # Re-scan paragraph after successful replacement

        if not replaced:
            # No more replaceable placeholders found
            break


def _replace_placeholders_in_table(
    table: Table,
    context: Dict[str, str],
    doc: DocxDocument = None
) -> None:
    """Replace all placeholders in a table.

    Note: Block-level content (lists) is not supported in table cells.

    Args:
        table: The table to process
        context: Dictionary mapping placeholder names to their values
        doc: The Word document (not used for tables, as block content not supported)
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Note: We don't pass doc to avoid inserting lists in table cells
                _replace_placeholders_in_paragraph(paragraph, context, doc=None)


def _replace_placeholders_in_document(doc: DocxDocument, context: Dict[str, str]) -> None:
    """Replace all placeholders in the entire document.

    Processes:
    - Main document body paragraphs
    - Tables in the main body
    - Headers and footers

    Args:
        doc: The Word document to process
        context: Dictionary mapping placeholder names to their values
    """
    # Process main body paragraphs
    for paragraph in doc.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, context, doc)

    # Process tables
    for table in doc.tables:
        _replace_placeholders_in_table(table, context, doc)

    # Process headers and footers
    for section in doc.sections:
        # Collect all header/footer parts to process
        parts = []

        # Default header and footer
        if section.header:
            parts.append(section.header)
        if section.footer:
            parts.append(section.footer)

        # First-page header/footer (when template uses "Different First Page")
        if section.different_first_page_header_footer:
            if section.first_page_header:
                parts.append(section.first_page_header)
            if section.first_page_footer:
                parts.append(section.first_page_footer)

        # Even-page header/footer (when template uses "Different Even & Odd Pages")
        even_page_header = getattr(section, 'even_page_header', None)
        even_page_footer = getattr(section, 'even_page_footer', None)
        if even_page_header:
            parts.append(even_page_header)
        if even_page_footer:
            parts.append(even_page_footer)

        for part in parts:
            for paragraph in part.paragraphs:
                # Headers/footers: don't support block content
                _replace_placeholders_in_paragraph(paragraph, context, doc=None)
            for table in part.tables:
                _replace_placeholders_in_table(table, context, doc=None)


def register_docx_template_tools_from_yaml(mcp: FastMCP, yaml_path: Path) -> None:
    """Register dynamic DOCX template tools from a YAML configuration file.

    Args:
        mcp: The FastMCP instance to register tools with
        yaml_path: Path to the YAML configuration file
    """
    try:
        cfg = yaml.safe_load(yaml_path.read_text(encoding="utf-8")) or {}
    except Exception as e:
        logger.error(f"[dynamic-docx] Failed to load YAML '{yaml_path}': {e}")
        return

    templates = cfg.get("templates") or []
    if not isinstance(templates, list):
        logger.error("[dynamic-docx] 'templates' key must be a list – skipping.")
        return

    for spec in templates:
        try:
            _register_single_template(mcp, spec)
        except Exception as e:
            name = spec.get("name", "<unknown>")
            logger.exception(f"[dynamic-docx] Failed to register template '{name}': {e}")


def _register_single_template(mcp: FastMCP, spec: Dict[str, Any]) -> None:
    """Register a single DOCX template as an MCP tool.

    Args:
        mcp: The FastMCP instance
        spec: The template specification from YAML
    """
    name = spec.get("name")
    if not name:
        logger.warning("[dynamic-docx] Template missing 'name', skipping.")
        return

    description = spec.get("description", f"Generate document from {name} template")
    annotations = spec.get("annotations", {})
    docx_path = spec.get("docx_path")

    if not docx_path:
        logger.warning(f"[dynamic-docx] Missing docx_path for {name}, skipping.")
        return

    # Validate path is filename only (no directory components)
    docx_path_obj = Path(docx_path)
    if docx_path_obj.is_absolute() or len(docx_path_obj.parts) != 1:
        logger.error(
            f"[dynamic-docx] docx_path must be filename only (no directories) for {name}; "
            f"got '{docx_path}'"
        )
        return

    # Resolve the template file
    resolved = find_docx_template_by_name(docx_path)
    if not resolved:
        logger.error(f"[dynamic-docx] Template file not found for {name}: {docx_path}")
        return

    logger.info(f"[dynamic-docx] Using template for {name}: {resolved}")

    # Build Pydantic model fields from args
    fields: Dict[str, Any] = {}

    for arg in spec.get("args", []):
        arg_name = arg.get("name")
        if not arg_name:
            continue

        # Handle enum values
        enum_values = arg.get("enum")
        if enum_values and isinstance(enum_values, list) and enum_values:
            if all(isinstance(v, int) for v in enum_values):
                lit_values = tuple(int(v) for v in enum_values)
            elif all(isinstance(v, (int, float)) for v in enum_values):
                lit_values = tuple(float(v) for v in enum_values)
            else:
                lit_values = tuple(str(v) for v in enum_values)
            py_type = Literal[lit_values]  # type: ignore[index]
            required = bool(arg.get("required", True))
            default = arg.get("default", (... if required else None))
            if default is not ... and default is not None and default not in lit_values:
                logger.warning(
                    f"[dynamic-docx] Default '{default}' not in enum for {arg_name}; ignoring default."
                )
                default = ... if required else None
            desc = arg.get("description") or f"One of: {', '.join(map(str, lit_values))}"
            fields[arg_name] = (py_type, Field(default, description=desc))
            continue

        # Handle regular types
        py_type = TYPE_MAP.get(str(arg.get("type", "string")).lower(), str)
        required = bool(arg.get("required", True))
        field_type = py_type if required else Optional[py_type]  # type: ignore[index]
        default = arg.get("default", (... if required else None))
        desc = arg.get("description", "")
        fields[arg_name] = (field_type, Field(default, description=desc) if desc else default)

    # Create the Pydantic model
    model = create_model(f"{name}_DocxArgs", **fields)  # type: ignore
    globals()[model.__name__] = model

    # Create the tool function
    def make_tool_fn(_model=model, _template_path=resolved, _name=name):
        def tool_impl(data: _model) -> str:  # type: ignore
            try:
                # Load the template document
                doc = DocxDocument(_template_path)

                # Build context from input data
                payload = data.model_dump()
                context = {k: ("" if v is None else str(v)) for k, v in payload.items()}

                # Replace placeholders
                _replace_placeholders_in_document(doc, context)

                # Save to buffer and upload
                buffer = io.BytesIO()
                try:
                    doc.save(buffer)
                    buffer.seek(0)

                    result = upload_file(buffer, "docx")
                finally:
                    buffer.close()

                logger.info(f"[dynamic-docx] Document generated from template {_name}")
                return result

            except Exception as e:
                logger.error(f"[dynamic-docx] Error generating document from {_name}: {e}", exc_info=True)
                raise ToolError(f"Error generating document from template {_name}: {e}")

        tool_impl.__annotations__['data'] = _model  # type: ignore[index]
        tool_impl.__annotations__['return'] = str  # type: ignore[index]
        return tool_impl

    # Register the tool
    mcp.tool(
        name=name,
        description=description,
        annotations=annotations,
        tags={"docx", "document", "template"},
    )(make_tool_fn())

    logger.info(f"[dynamic-docx] Registered tool: {name}")

