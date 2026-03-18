import logging
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from template_utils import find_docx_template

logger = logging.getLogger(__name__)


def load_templates():
    """Resolve Word template path from custom/default template directories.

    Returns absolute path as string or None if not found.
    """
    path = find_docx_template()
    if path:
        logger.debug(f"Using Word template: {path}")
    else:
        logger.warning("No Word template found, will create a blank document")
    return path


def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    """Adds a hyperlink to a paragraph.

    Falls back to plain text if hyperlink creation fails.
    """
    try:
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)

        new_run.append(rPr)

        text_elem = OxmlElement('w:t')
        text_elem.text = text
        text_elem.set(qn('xml:space'), 'preserve')
        new_run.append(text_elem)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception as e:
        logger.warning("Failed to create hyperlink for '%s' (%s), falling back to plain text: %s", text, url, e)
        paragraph.add_run(text)


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

def parse_inline_formatting(text, paragraph, bold=False, italic=False):
    """Parse inline markdown formatting like **bold**, *italic*, and [links](url)

    Args:
        text: The text to parse
        paragraph: The paragraph to add runs to
        bold: Whether the current context is bold (for nested formatting)
        italic: Whether the current context is italic (for nested formatting)
    """
    escape_ctx = {"map": {}, "counter": 0}
    text = _handle_escapes(text, escape_ctx)

    line_parts = text.split('  \n')
    for line_idx, line_part in enumerate(line_parts):
        if not line_part and line_idx == len(line_parts) - 1:
            continue
        _parse_formatting_segment(line_part, paragraph, bold, italic, escape_ctx)
        if line_idx < len(line_parts) - 1:
            paragraph.add_run().add_break()


def _apply_formatting(run, bold=False, italic=False):
    """Apply inherited bold/italic formatting to a run."""
    if bold:
        run.bold = True
    if italic:
        run.italic = True


_INLINE_FORMAT_RE = re.compile(
    r'(\*{3}(?:[^*]|\*(?!\*{2}))+\*{3}'  # ***bold italic***
    r'|\*\*(?:[^*]|\*(?!\*))+\*\*'       # **bold**
    r'|~~.+?~~'                           # ~~strikethrough~~
    r'|__(?!_).+?__'                      # __underline__
    r'|\*(?:[^*]|\*\*[^*]+\*\*)+\*'       # *italic* (allows nested **bold**)
    r'|`[^`]+`'                           # `code`
    r'|\[[^\]]*\]\([^)]*\))'             # [link](url)
)


def _parse_formatting_segment(text, paragraph, bold=False, italic=False, escape_ctx=None):
    """Parse a single text segment for inline markdown formatting."""
    for part in _INLINE_FORMAT_RE.split(text):
        if not part:
            continue
        if part.startswith('***') and part.endswith('***') and len(part) > 6:
            _parse_formatting_segment(part[3:-3], paragraph, bold=True, italic=True, escape_ctx=escape_ctx)
        elif part.startswith('**') and part.endswith('**'):
            _parse_formatting_segment(part[2:-2], paragraph, bold=True, italic=italic, escape_ctx=escape_ctx)
        elif part.startswith('~~') and part.endswith('~~'):
            run = paragraph.add_run(_restore_escapes(part[2:-2], escape_ctx))
            run.font.strike = True
            _apply_formatting(run, bold, italic)
        elif part.startswith('__') and part.endswith('__') and not part.startswith('___'):
            run = paragraph.add_run(_restore_escapes(part[2:-2], escape_ctx))
            run.font.underline = True
            _apply_formatting(run, bold, italic)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            _parse_formatting_segment(part[1:-1], paragraph, bold=bold, italic=True, escape_ctx=escape_ctx)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(_restore_escapes(part[1:-1], escape_ctx))
            run.font.name = 'Courier New'
            _apply_formatting(run, bold, italic)
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            link_match = re.match(r'\[(.*?)]\((.*?)\)', part)
            if link_match:
                link_text = _restore_escapes(link_match.group(1), escape_ctx)
                link_url = _restore_escapes(link_match.group(2), escape_ctx)
                add_hyperlink(paragraph, link_text, link_url)
        else:
            _apply_formatting(paragraph.add_run(_restore_escapes(part, escape_ctx)), bold, italic)


def _handle_escapes(text, escape_ctx):
    """Replace backslash-escaped characters with PUA placeholders.

    The placeholders survive through the formatting regex so that escaped
    characters (e.g. ``\\*``) are **not** treated as markdown markers.
    Call :func:`_restore_escapes` on final text before inserting into runs.
    """

    def _replace(match):
        placeholder = chr(0xE000 + escape_ctx["counter"])
        escape_ctx["map"][placeholder] = match.group(1)
        escape_ctx["counter"] += 1
        return placeholder

    return re.sub(r'\\(.)', _replace, text)


def _restore_escapes(text, escape_ctx):
    """Replace PUA placeholders back with their original literal characters."""
    esc_map = escape_ctx["map"] if escape_ctx else {}
    if not esc_map:
        return text
    for placeholder, char in esc_map.items():
        text = text.replace(placeholder, char)
    return text


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def parse_table(lines, start_idx):
    """Parse markdown table and return the table data and next line index."""
    table_lines = []
    i = start_idx

    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
        else:
            break

    if len(table_lines) < 2:
        return None, start_idx + 1

    table_data = []
    for line in table_lines:
        if '---' in line or ':-:' in line or ':--' in line or '--:' in line:
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        table_data.append(cells)

    return table_data, i


def add_table_to_doc(table_data, doc):
    """Add table data to Word document.

    Returns the created ``Table`` object, or ``None`` when the table could
    not be created (empty data or exception).
    """
    if not table_data:
        return None

    rows = len(table_data)
    cols = max(len(row) for row in table_data) if table_data else 0

    try:
        word_table = doc.add_table(rows=rows, cols=cols)
        word_table.style = 'Table Grid'
    except Exception as e:
        logger.warning("Failed to create table with 'Table Grid' style, using default: %s", e)
        try:
            word_table = doc.add_table(rows=rows, cols=cols)
        except Exception as e2:
            logger.error("Failed to create table: %s", e2, exc_info=True)
            return None

    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                try:
                    cell = word_table.cell(i, j)
                    if cell.paragraphs:
                        cell.paragraphs[0].clear()
                    parse_inline_formatting(cell_text, cell.paragraphs[0])
                except Exception as e:
                    logger.warning("Failed to populate table cell [%d, %d]: %s", i, j, e)

    return word_table


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------

def process_list_items(lines, start_idx, doc, is_ordered=False, level=0,
                       return_elements=False):
    """Process markdown list items with proper Word numbering.

    When *return_elements* is True the created paragraph XML elements are
    removed from the document body and returned so the caller can re-insert
    them elsewhere (used by the template placeholder machinery).

    Returns:
        Tuple of (next_line_index, list_of_elements | None).
    """
    bullet_styles = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
    number_styles = ['List Number', 'List Number 2', 'List Number 3']

    style_array = number_styles if is_ordered else bullet_styles
    style = style_array[min(level, len(style_array) - 1)]

    elements = [] if return_elements else None
    i = start_idx

    while i < len(lines):
        line = lines[i].strip()

        original_line = lines[i]
        indent = len(original_line) - len(original_line.lstrip())
        current_level = indent // 3

        if current_level != level:
            break

        if is_ordered:
            list_match = re.match(r'^\d+\.\s+(.+)', line)
        else:
            list_match = re.match(r'^[-*+]\s+(.+)', line)

        if not list_match:
            break

        paragraph = doc.add_paragraph(style=style)
        parse_inline_formatting(list_match.group(1), paragraph)

        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)

        i += 1

        # Look ahead for nested items
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                i += 1
                continue

            next_original = lines[i]
            next_indent = len(next_original) - len(next_original.lstrip())
            next_level = next_indent // 3

            if next_level > level:
                is_nested_ordered = bool(re.match(r'^\d+\.\s+', next_line))
                is_nested_unordered = bool(re.match(r'^[-*+]\s+', next_line))
                if is_nested_ordered or is_nested_unordered:
                    i, nested = process_list_items(
                        lines, i, doc, is_nested_ordered, next_level, return_elements
                    )
                    if return_elements and nested:
                        elements.extend(nested)
                else:
                    break
            else:
                break

    return i, elements


# ---------------------------------------------------------------------------
# Block-level patterns (compiled once, used by many modules)
# ---------------------------------------------------------------------------

ORDERED_LIST_PATTERN = re.compile(r'^\d+\.\s+')
UNORDERED_LIST_PATTERN = re.compile(r'^[-*+]\s+')
HEADING_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$')
PAGE_BREAK_PATTERN = re.compile(r'^-{3,}\s*$')
HORIZONTAL_LINE_PATTERN = re.compile(r'^\*{3,}\s*$')
IMAGE_PATTERN = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)$')
TABLE_LINE_PATTERN = re.compile(r'^\|.+\|$')

# All block-level patterns checked by contains_block_markdown
_BLOCK_PATTERNS = [
    ORDERED_LIST_PATTERN, UNORDERED_LIST_PATTERN, HEADING_PATTERN,
    PAGE_BREAK_PATTERN, HORIZONTAL_LINE_PATTERN, IMAGE_PATTERN,
    TABLE_LINE_PATTERN,
]


def contains_block_markdown(value: str) -> bool:
    """Return True if *value* contains block-level markdown content."""
    for line in value.split('\n'):
        stripped = line.strip()
        if any(p.match(stripped) for p in _BLOCK_PATTERNS):
            return True
        if detect_alignment(stripped) is not None:
            return True
    return False


# ---------------------------------------------------------------------------
# Page break / horizontal line
# ---------------------------------------------------------------------------

def add_horizontal_line(doc):
    """Add a visual horizontal line (thin border) to the document."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------

def add_image_to_doc(doc, url, alt_text, max_width_inches=None):
    """Add an image from a URL to the document.

    Downloads the image and inserts it.  On failure inserts an error
    placeholder paragraph instead.
    """
    try:
        from pptx_tools.image_utils import download_image

        if max_width_inches is None:
            try:
                sec = doc.sections[-1]
                max_width_inches = (sec.page_width - sec.left_margin - sec.right_margin) / 914400
            except Exception:
                max_width_inches = 5.5

        image_stream, _ = download_image(url)
        doc.add_picture(image_stream, width=Inches(max_width_inches))

        if alt_text:
            caption = doc.add_paragraph()
            caption.add_run(alt_text).italic = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning("Failed to add image from '%s': %s", url, e)
        doc.add_paragraph().add_run(f"[Image could not be loaded: {url}]")


# ---------------------------------------------------------------------------
# Text alignment
# ---------------------------------------------------------------------------

ALIGNMENT_MAP = {
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
}

# Inline (single-line):  <center>text</center>  or  <div align="x">text</div>
_ALIGN_INLINE_RE = re.compile(
    r'^(?:<center>(.*)</center>'
    r'|<div\s+align="(right|center|justify|left)">(.*)</div>)$',
    re.IGNORECASE,
)
# Block open:  <center>  or  <div align="x">  (content on following lines)
_ALIGN_OPEN_RE = re.compile(
    r'^(?:<center>'
    r'|<div\s+align="(right|center|justify|left)">)\s*$',
    re.IGNORECASE,
)
# Block close:  </center>  or  </div>
_ALIGN_CLOSE_RE = re.compile(r'^</(?:center|div)>\s*$', re.IGNORECASE)


def detect_alignment(line):
    """Detect an alignment tag (inline *or* block-open) on *line*.

    Returns ``(inner_text, alignment)`` for an inline tag,
    ``(None, alignment)`` for a block-open tag, or ``None`` if no match.
    """
    m = _ALIGN_INLINE_RE.match(line)
    if m:
        # group(1) = center content, group(2) = div-align value, group(3) = div content
        if m.group(1) is not None:
            return m.group(1).strip(), WD_ALIGN_PARAGRAPH.CENTER
        return m.group(3).strip(), ALIGNMENT_MAP.get(m.group(2).lower(), WD_ALIGN_PARAGRAPH.LEFT)

    m = _ALIGN_OPEN_RE.match(line)
    if m:
        align = ALIGNMENT_MAP.get((m.group(1) or 'center').lower(), WD_ALIGN_PARAGRAPH.CENTER)
        return None, align

    return None


def process_alignment_block(lines, start_idx, doc, alignment, return_elements=False):
    """Process lines inside a multi-line alignment block."""
    elements = [] if return_elements else None
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if _ALIGN_CLOSE_RE.match(stripped):
            i += 1
            break
        if not stripped:
            i += 1
            continue
        para = doc.add_paragraph()
        para.alignment = alignment
        parse_inline_formatting(stripped, para)
        if return_elements:
            elements.append(para._p)
            doc._body._body.remove(para._p)
        i += 1
    return i, elements


# ---------------------------------------------------------------------------
# Header / footer / Word fields
# ---------------------------------------------------------------------------

def _add_field(paragraph, field_code):
    """Insert a Word field (PAGE, NUMPAGES, etc.) into a paragraph."""
    for fld_type, text in [('begin', None), (None, field_code), ('end', None)]:
        run = paragraph.add_run()
        if fld_type:
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), fld_type)
            run._r.append(fld)
        else:
            elem = OxmlElement('w:instrText')
            elem.set(qn('xml:space'), 'preserve')
            elem.text = f' {text} '
            run._r.append(elem)


_PAGE_TOKEN_RE = re.compile(r'(\{page}|\{pages})')


def set_header_footer(doc, text, kind='header'):
    """Set document header or footer text.

    Iterates over **all** document sections.  For each section the default
    header/footer is updated, and — when the section uses a different first-page
    header/footer — that variant is updated as well.

    Pre-existing paragraph formatting (alignment, style) from the template is
    preserved; only run content is replaced.

    Args:
        doc: The Word document.
        text: Content string.  Use ``{page}`` / ``{pages}`` for field tokens.
        kind: ``'header'`` or ``'footer'``.
    """
    _TOKEN_MAP = {'{page}': 'PAGE', '{pages}': 'NUMPAGES'}

    def _fill_paragraph(p, content):
        """Clear existing runs/fields and write *content* into paragraph *p*."""
        # Preserve existing alignment if set
        existing_alignment = p.alignment

        # Remove all existing child run (<w:r>) and field elements
        for child in list(p._p):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('r', 'hyperlink', 'fldSimple'):
                p._p.remove(child)

        for part in _PAGE_TOKEN_RE.split(content):
            if part in _TOKEN_MAP:
                _add_field(p, _TOKEN_MAP[part])
            elif part:
                p.add_run(part)

        # Restore alignment – fall back to CENTER when the template had none
        p.alignment = existing_alignment if existing_alignment is not None else WD_ALIGN_PARAGRAPH.CENTER

    def _update_part(section_part):
        """Update a single header or footer part."""
        section_part.is_linked_to_previous = False
        if section_part.paragraphs:
            _fill_paragraph(section_part.paragraphs[0], text)
        else:
            p = section_part.add_paragraph()
            _fill_paragraph(p, text)

    for section in doc.sections:
        # Default header / footer
        _update_part(getattr(section, kind))

        # First-page header / footer (when the template enables it)
        if section.different_first_page_header_footer:
            first_kind = f'first_page_{kind}'
            first_part = getattr(section, first_kind, None)
            if first_part is not None:
                _update_part(first_part)

        # Even-page header / footer
        even_kind = f'even_page_{kind}'
        even_part = getattr(section, even_kind, None)
        if even_part is not None and doc.settings.element.find(qn('w:evenAndOddHeaders')) is not None:
            _update_part(even_part)


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

def add_toc(doc):
    """Insert a Table of Contents field.

    The TOC is based on Heading styles 1-3 and will update when the document
    is opened in Word.
    """
    doc.add_heading('Table of Contents', level=1)

    p = doc.add_paragraph()

    # begin field
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld)

    # instruction
    run = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    # separate
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld)

    # placeholder text
    p.add_run('[Table of Contents — open in Word and press F9 to update]')

    # end field
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'end')
    run._r.append(fld)

    doc.add_page_break()

    # Tell Word to update fields on open
    uf = OxmlElement('w:updateFields')
    uf.set(qn('w:val'), 'true')
    doc.settings.element.append(uf)


# ---------------------------------------------------------------------------
# Generic block processor (used by dynamic template placeholder replacement)
# ---------------------------------------------------------------------------

def process_markdown_block(doc, lines, start_idx, return_element=True):
    """Process a single markdown block element and return created XML elements.

    Returns:
        Tuple of (next_index, list_of_elements).
    """
    line = lines[start_idx]
    stripped = line.strip()
    elements = []

    def _collect(element):
        """If return_element, detach *element* (paragraph or table) from body and collect it."""
        if return_element:
            elements.append(element)
            doc._body._body.remove(element)

    try:
        # Heading
        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading = doc.add_heading('', level=min(level, 6))
            parse_inline_formatting(heading_match.group(2), heading)
            _collect(heading._p)
            return start_idx + 1, elements

        # Table (lines starting with |)
        if TABLE_LINE_PATTERN.match(stripped):
            table_data, next_idx = parse_table(lines, start_idx)
            if table_data:
                word_table = add_table_to_doc(table_data, doc)
                if word_table is not None:
                    _collect(word_table._tbl)
                return next_idx, elements
            # Not a valid table (e.g. single pipe line) — fall through to
            # emit as a regular paragraph below.

        # Page break (---)
        if PAGE_BREAK_PATTERN.match(stripped):
            doc.add_page_break()
            _collect(doc.paragraphs[-1]._p)
            return start_idx + 1, elements

        # Horizontal line (***)
        if HORIZONTAL_LINE_PATTERN.match(stripped):
            _collect(add_horizontal_line(doc)._p)
            return start_idx + 1, elements

        # Image (![alt](url))
        img_match = IMAGE_PATTERN.match(stripped)
        if img_match:
            add_image_to_doc(doc, img_match.group(2), img_match.group(1))
            return start_idx + 1, elements

        # Alignment (inline or block-open)
        align_result = detect_alignment(stripped)
        if align_result is not None:
            inner, alignment = align_result
            if inner is not None:
                # Single-line
                para = doc.add_paragraph()
                para.alignment = alignment
                parse_inline_formatting(inner, para)
                _collect(para._p)
                return start_idx + 1, elements
            else:
                # Multi-line block
                idx, block_elems = process_alignment_block(
                    lines, start_idx + 1, doc, alignment, return_elements=return_element
                )
                if return_element and block_elems:
                    elements.extend(block_elems)
                return idx, elements

        # Ordered list
        if ORDERED_LIST_PATTERN.match(stripped):
            return process_list_items(
                lines, start_idx, doc, is_ordered=True, level=0, return_elements=return_element
            )

        # Unordered list
        if UNORDERED_LIST_PATTERN.match(stripped):
            return process_list_items(
                lines, start_idx, doc, is_ordered=False, level=0, return_elements=return_element
            )

        # Regular paragraph
        para = doc.add_paragraph()
        parse_inline_formatting(stripped, para)
        _collect(para._p)
        return start_idx + 1, elements

    except Exception as e:
        logger.error("Failed to process markdown block at line %d: %s", start_idx, e, exc_info=True)
        return start_idx + 1, elements
