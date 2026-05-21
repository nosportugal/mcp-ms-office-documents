import io
import logging
from docx import Document

from upload_tools import upload_file
from .helpers import (
    load_templates,
    parse_inline_formatting,
    process_markdown_block,
    set_header_footer,
    add_toc,
)

logger = logging.getLogger(__name__)


def _markdown_to_doc(markdown_content, title=None, author=None, subject=None,
                     header_text=None, footer_text=None, include_toc=False):
    """Convert Markdown content to a python-docx Document object.

    This is the core conversion logic, separated from upload concerns so it
    can be used directly in tests or other contexts that need the Document.

    Returns:
        A ``docx.Document`` instance with the rendered content.
    """
    logger.info("Starting markdown_to_doc conversion")
    path = load_templates()

    # Create document with or without template
    try:
        if path:
            logger.debug(f"Using Word template at: {path}")
            doc = Document(path)
        else:
            doc = Document()  # Create blank document if no template
            logger.warning("No template found, creating blank document")
    except Exception as e:
        logger.error("Failed to load Word template '%s': %s", path, e, exc_info=True)
        raise RuntimeError(f"Error loading Word template: {e}") from e

    # Set document metadata
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author
    if subject:
        doc.core_properties.subject = subject

    # Insert Table of Contents if requested (before main content)
    if include_toc:
        add_toc(doc)

    # Set header and footer
    if header_text:
        set_header_footer(doc, header_text, 'header')
    if footer_text:
        set_header_footer(doc, footer_text, 'footer')

    # Parse markdown content into document
    lines = markdown_content.split('\n')
    n = len(lines)
    i = 0
    blocks_count = 0

    try:
        while i < n:
            line = lines[i]

            # --- Empty line handling (preserve spacing) ---
            if not line.strip():
                empty_line_count = 1
                i += 1

                while i < n and not lines[i].strip():
                    empty_line_count += 1
                    i += 1

                # Two or more consecutive empty lines → insert spacing paragraphs
                if empty_line_count >= 2:
                    for _ in range(empty_line_count - 1):
                        doc.add_paragraph()
                continue

            # --- Soft line breaks (trailing two spaces) ---
            # Markdown convention: a line ending with two spaces means a soft
            # break within the same paragraph. Collect consecutive lines that
            # are joined by this mechanism.
            if line.endswith('  '):
                paragraph_lines = []
                while i < n:
                    current_line = lines[i]
                    if not current_line.strip():
                        break
                    paragraph_lines.append(current_line)
                    i += 1
                    if not current_line.endswith('  '):
                        break

                full_text = '  \n'.join(paragraph_lines)
                first_line = paragraph_lines[0].strip()

                if first_line.startswith('#'):
                    stripped_hashes = first_line.lstrip('#')
                    level = len(first_line) - len(stripped_hashes)
                    heading = doc.add_heading('', level=min(level, 6))
                    parse_inline_formatting(stripped_hashes.strip(), heading)
                elif first_line.startswith('>'):
                    quote_text = full_text[1:].strip()
                    quote_paragraph = doc.add_paragraph()
                    quote_paragraph.style = 'Quote'
                    parse_inline_formatting(quote_text, quote_paragraph)
                else:
                    paragraph = doc.add_paragraph()
                    parse_inline_formatting(full_text, paragraph)

                blocks_count += 1
                continue

            # --- All other block elements: delegate to shared processor ---
            i, _ = process_markdown_block(doc, lines, i, return_element=False)
            blocks_count += 1

    except Exception as e:
        logger.error(f"Error in parsing markdown: {e}", exc_info=True)
        raise RuntimeError(f"Error in parsing markdown: {e}") from e

    logger.info(f"Markdown parsing completed ({blocks_count} blocks processed)")
    return doc


def markdown_to_word(markdown_content, title=None, author=None, subject=None,
                     header_text=None, footer_text=None, include_toc=False, file_name=None):
    """Convert Markdown to Word document, save to memory and upload."""
    doc = _markdown_to_doc(
        markdown_content,
        title=title,
        author=author,
        subject=subject,
        header_text=header_text,
        footer_text=footer_text,
        include_toc=include_toc,
    )

    # Save the document to BytesIO and upload
    try:
        logger.info("Saving Word document to memory buffer")
        file_object = io.BytesIO()
        doc.save(file_object)
        file_object.seek(0)

        result = upload_file(file_object, "docx", filename=file_name)
        file_object.close()

        logger.info("Word document uploaded successfully")
        return result
    except Exception as e:
        logger.error(f"Error saving/uploading Word document: {e}", exc_info=True)
        raise RuntimeError(f"Error saving/uploading Word document: {e}") from e
