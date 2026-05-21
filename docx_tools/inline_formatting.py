"""Inline markdown formatting: bold, italic, strikethrough, underline, code, links.
Handles escape sequences and hyperlink insertion into python-docx paragraphs.
"""
import html
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from .patterns import _INLINE_FORMAT_RE, _LINK_RE, _ESCAPE_RE, _BR_RE

logger = logging.getLogger(__name__)

# HTML entities safe to decode (no markdown/alignment meaning).
# Order: most common first for early short-circuit on the `in` check.
_SAFE_HTML_ENTITIES = (
    ('&nbsp;', '\u00a0'),     # non-breaking space
    ('&ndash;', '\u2013'),    # en-dash
    ('&mdash;', '\u2014'),    # em-dash
    ('&hellip;', '\u2026'),   # ellipsis …
    ('&bull;', '\u2022'),     # bullet •
    ('&trade;', '\u2122'),    # ™
    ('&copy;', '\u00a9'),     # ©
    ('&reg;', '\u00ae'),      # ®
    ('&deg;', '\u00b0'),      # °
    ('&plusmn;', '\u00b1'),   # ±
    ('&times;', '\u00d7'),    # ×
    ('&divide;', '\u00f7'),   # ÷
    ('&lsquo;', '\u2018'),   # '
    ('&rsquo;', '\u2019'),   # '
    ('&ldquo;', '\u201c'),   # "
    ('&rdquo;', '\u201d'),   # "
    ('&euro;', '\u20ac'),     # €
)
def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    """Adds a hyperlink to a paragraph.
    Falls back to plain text if hyperlink creation fails.
    """
    try:
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)
        new_run.append(rPr)
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        text_elem.set(qn('xml:space'), 'preserve')
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception as e:
        logger.warning("Failed to create hyperlink for '%s' (%s), falling back to plain text: %s", text, url, e)
        paragraph.add_run(text)
def _handle_escapes(text, escape_ctx):
    """Replace backslash-escaped characters with PUA placeholders."""
    def _replace(match):
        placeholder = chr(0xE000 + escape_ctx["counter"])
        escape_ctx["map"][placeholder] = match.group(1)
        escape_ctx["counter"] += 1
        return placeholder
    return _ESCAPE_RE.sub(_replace, text)
def _restore_escapes(text, escape_ctx):
    """Replace PUA placeholders back with their original literal characters."""
    esc_map = escape_ctx["map"] if escape_ctx else {}
    if not esc_map:
        return text
    for placeholder, char in esc_map.items():
        text = text.replace(placeholder, char)
    return text
def _apply_formatting(run, bold=False, italic=False):
    """Apply inherited bold/italic formatting to a run."""
    if bold:
        run.bold = True
    if italic:
        run.italic = True
def parse_inline_formatting(text, paragraph, bold=False, italic=False):
    """Parse inline markdown formatting like **bold**, *italic*, and [links](url)
    Args:
        text: The text to parse
        paragraph: The paragraph to add runs to
        bold: Whether the current context is bold (for nested formatting)
        italic: Whether the current context is italic (for nested formatting)
    """
    # Decode common HTML entities that LLMs emit and that are safe (no markdown
    # meaning).  We intentionally do NOT use html.unescape() because decoding
    # entities like &lt; &gt; &amp; &#42; etc. would produce characters that the
    # downstream markdown parser would misinterpret as formatting markers.
    for entity, char in _SAFE_HTML_ENTITIES:
        if entity in text:
            text = text.replace(entity, char)
    # Normalize <br>, <br/>, <br /> tags to the two-space soft-break marker
    # so they produce line breaks (common in table cells where real newlines
    # would break the row).
    text = _BR_RE.sub('  \n', text)
    escape_ctx = {"map": {}, "counter": 0}
    text = _handle_escapes(text, escape_ctx)
    line_parts = text.split('  \n')
    for line_idx, line_part in enumerate(line_parts):
        if not line_part and line_idx == len(line_parts) - 1:
            continue
        _parse_formatting_segment(line_part, paragraph, bold, italic, escape_ctx)
        if line_idx < len(line_parts) - 1:
            paragraph.add_run().add_break()
def _parse_formatting_segment(text, paragraph, bold=False, italic=False, escape_ctx=None):
    """Parse a single text segment for inline markdown formatting."""
    for part in _INLINE_FORMAT_RE.split(text):
        if not part:
            continue
        if part.startswith('***') and part.endswith('***') and len(part) > 6:
            _parse_formatting_segment(part[3:-3], paragraph, bold=True, italic=True, escape_ctx=escape_ctx)
        elif part.startswith('**') and part.endswith('**'):
            _parse_formatting_segment(part[2:-2], paragraph, bold=True, italic=italic, escape_ctx=escape_ctx)
        elif part.startswith('~~') and part.endswith('~~'):
            run = paragraph.add_run(_restore_escapes(part[2:-2], escape_ctx))
            run.font.strike = True
            _apply_formatting(run, bold, italic)
        elif part.startswith('==') and part.endswith('==') and len(part) > 4:
            from docx.enum.text import WD_COLOR_INDEX
            run = paragraph.add_run(_restore_escapes(part[2:-2], escape_ctx))
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            _apply_formatting(run, bold, italic)
        elif part.startswith('__') and part.endswith('__') and not part.startswith('___'):
            run = paragraph.add_run(_restore_escapes(part[2:-2], escape_ctx))
            run.font.underline = True
            _apply_formatting(run, bold, italic)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            _parse_formatting_segment(part[1:-1], paragraph, bold=bold, italic=True, escape_ctx=escape_ctx)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(_restore_escapes(part[1:-1], escape_ctx))
            run.font.name = 'Courier New'
            _apply_formatting(run, bold, italic)
        elif part.startswith('^') and part.endswith('^') and len(part) > 2:
            run = paragraph.add_run(_restore_escapes(part[1:-1], escape_ctx))
            run.font.superscript = True
            _apply_formatting(run, bold, italic)
        elif part.startswith('~') and part.endswith('~') and not part.startswith('~~') and len(part) > 2:
            run = paragraph.add_run(_restore_escapes(part[1:-1], escape_ctx))
            run.font.subscript = True
            _apply_formatting(run, bold, italic)
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            link_match = _LINK_RE.match(part)
            if link_match:
                link_text = _restore_escapes(link_match.group(1), escape_ctx)
                link_url = _restore_escapes(link_match.group(2), escape_ctx)
                add_hyperlink(paragraph, link_text, link_url)
        else:
            _apply_formatting(paragraph.add_run(_restore_escapes(part, escape_ctx)), bold, italic)
