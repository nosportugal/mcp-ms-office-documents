"""Tests for base DOCX tool (markdown_to_word).

These tests verify that the markdown to Word conversion works correctly,
including headers, lists, tables, formatting, links, and block quotes.

Output files are saved to tests/output/docx/ directory for manual inspection.
"""

import sys
from pathlib import Path

# Add project root to path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_tools.base_docx_tool import _markdown_to_doc
from docx_tools.inline_formatting import parse_inline_formatting

# Output directory for test files
OUTPUT_DIR = Path(__file__).parent / "output" / "docx"


@pytest.fixture(scope="module", autouse=True)
def setup_output_dir():
    """Create output directory if it doesn't exist."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    yield


def create_word_document(markdown_content: str, title=None, author=None,
                         subject=None, header_text=None, footer_text=None,
                         include_toc=False) -> Document:
    """Convert Markdown to Word document and return the Document object.

    Delegates to the production _markdown_to_doc function so tests exercise
    the real code path rather than a divergent copy.
    """
    return _markdown_to_doc(
        markdown_content,
        title=title,
        author=author,
        subject=subject,
        header_text=header_text,
        footer_text=footer_text,
        include_toc=include_toc,
    )


def save_test_document(markdown: str, filename: str) -> Document:
    """Convert markdown to Word and save directly to test output directory.

    Args:
        markdown: Markdown content to convert
        filename: Output filename (e.g., 'header_h1.docx')

    Returns:
        The generated Document object for assertions
    """
    doc = create_word_document(markdown)
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    return doc


# =============================================================================
# Header Tests
# =============================================================================

class TestHeaders:
    """Tests for markdown headers conversion."""

    def test_h1_header(self):
        """Test H1 header conversion."""
        markdown = "# Main Title"
        doc = save_test_document(markdown, "header_h1.docx")
        assert doc is not None

    def test_h2_header(self):
        """Test H2 header conversion."""
        markdown = "## Section Title"
        doc = save_test_document(markdown, "header_h2.docx")
        assert doc is not None

    def test_h3_header(self):
        """Test H3 header conversion."""
        markdown = "### Subsection Title"
        doc = save_test_document(markdown, "header_h3.docx")
        assert doc is not None

    def test_multiple_headers(self):
        """Test document with multiple header levels."""
        markdown = """# Document Title

## Introduction

Some intro text here.

### Details

More details.

## Conclusion

Final thoughts.
"""
        doc = save_test_document(markdown, "header_multiple.docx")
        assert doc is not None

    def test_header_with_formatting(self):
        """Test header with inline formatting."""
        markdown = "# Title with **bold** and *italic*"
        doc = save_test_document(markdown, "header_formatted.docx")
        assert doc is not None


# =============================================================================
# List Tests
# =============================================================================

class TestLists:
    """Tests for markdown list conversion."""

    def test_unordered_list(self):
        """Test unordered (bullet) list."""
        markdown = """- First item
- Second item
- Third item
"""
        doc = save_test_document(markdown, "list_unordered.docx")
        assert doc is not None

    def test_ordered_list(self):
        """Test ordered (numbered) list."""
        markdown = """1. First item
2. Second item
3. Third item
"""
        doc = save_test_document(markdown, "list_ordered.docx")
        assert doc is not None

    def test_nested_list(self):
        """Test nested list items."""
        markdown = """- Main item 1
   - Sub item 1.1
   - Sub item 1.2
- Main item 2
   - Sub item 2.1
"""
        doc = save_test_document(markdown, "list_nested.docx")
        assert doc is not None

    def test_list_with_formatting(self):
        """Test list items with inline formatting."""
        markdown = """- **Bold item**
- *Italic item*
- Item with `code`
- Item with [link](https://example.com)
"""
        doc = save_test_document(markdown, "list_formatted.docx")
        assert doc is not None

    def test_mixed_list_types(self):
        """Test document with both ordered and unordered lists."""
        markdown = """## Shopping List

- Apples
- Bananas
- Oranges

## Steps to Follow

1. First step
2. Second step
3. Third step
"""
        doc = save_test_document(markdown, "list_mixed.docx")
        assert doc is not None


# =============================================================================
# Table Tests
# =============================================================================

class TestTables:
    """Tests for markdown table conversion."""

    def test_simple_table(self):
        """Test simple table conversion."""
        markdown = """| Name | Age | City |
|------|-----|------|
| John | 25  | NYC  |
| Jane | 30  | LA   |
"""
        doc = save_test_document(markdown, "table_simple.docx")
        assert doc is not None

    def test_table_with_formatting(self):
        """Test table with formatted cells."""
        markdown = """| Feature | Description |
|---------|-------------|
| **Bold** | This is bold |
| *Italic* | This is italic |
| `Code` | This is code |
"""
        doc = save_test_document(markdown, "table_formatted.docx")
        assert doc is not None

    def test_table_with_alignment(self):
        """Test table with column alignment markers applied to cells."""
        markdown = """| Left | Center | Right |
|:-----|:------:|------:|
| L1   | C1     | R1    |
| L2   | C2     | R2    |
"""
        doc = save_test_document(markdown, "table_aligned.docx")
        assert doc is not None
        table = doc.tables[0]
        # Header row + 2 data rows
        assert len(table.rows) == 3
        # Check alignment on data cells
        # Left column (:---) — default, so alignment is None
        assert table.cell(1, 0).paragraphs[0].alignment is None
        # Center column (:---:)
        assert table.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        # Right column (---:)
        assert table.cell(1, 2).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_table_borderless(self):
        """Test that <!-- borderless --> directive removes table borders."""
        markdown = """<!-- borderless -->
| English | French |
|---------|--------|
| Hello   | Bonjour |
| Goodbye | Au revoir |
"""
        doc = save_test_document(markdown, "table_borderless.docx")
        assert doc is not None
        table = doc.tables[0]
        # Verify borders are set to 'none'
        tblPr = table._tbl.tblPr
        borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        assert borders is not None
        # Check that all border elements have val='none'
        for border in borders:
            assert border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') == 'none'
        # Verify the directive itself didn't create a paragraph
        # (only the table should be in the doc body, no "<!-- borderless -->" text)
        for para in doc.paragraphs:
            assert '<!--' not in para.text

    def test_table_column_widths(self):
        """Test that <!-- widths: ... --> sets proportional column widths."""
        markdown = """<!-- widths: 30 70 -->
| Narrow | Wide |
|--------|------|
| A      | This column should be wider |
"""
        doc = save_test_document(markdown, "table_col_widths_2col.docx")
        assert doc is not None
        table = doc.tables[0]
        # Column 0 should be narrower than column 1
        col0_width = table.cell(0, 0).width
        col1_width = table.cell(0, 1).width
        assert col1_width > col0_width

    def test_table_column_widths_3col(self):
        """Test column widths with 3 columns."""
        markdown = """<!-- widths: 20 50 30 -->
| Small | Large | Medium |
|-------|-------|--------|
| A     | B     | C      |
"""
        doc = save_test_document(markdown, "table_col_widths_3col.docx")
        assert doc is not None
        table = doc.tables[0]
        col0_width = table.cell(0, 0).width
        col1_width = table.cell(0, 1).width
        col2_width = table.cell(0, 2).width
        # col1 (50) > col2 (30) > col0 (20)
        assert col1_width > col2_width
        assert col2_width > col0_width

    def test_table_combined_directives(self):
        """Test borderless + widths directives together."""
        markdown = """<!-- borderless -->
<!-- widths: 40 60 -->
| English | French |
|---------|--------|
| Hello   | Bonjour |
| Goodbye | Au revoir |
"""
        doc = save_test_document(markdown, "table_combined_directives.docx")
        assert doc is not None
        table = doc.tables[0]
        # Verify borderless
        tblPr = table._tbl.tblPr
        borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        assert borders is not None
        # Verify widths applied (col1 wider than col0)
        assert table.cell(0, 1).width > table.cell(0, 0).width

    def test_table_cell_line_breaks(self):
        """Test that <br> in table cells creates multiple paragraphs."""
        markdown = """| Header 1 | Header 2 |
|----------|----------|
| Line one<br>Line two | Single line |
| **Bold**<br>*Italic* | A<br/>B<br>C |
"""
        doc = save_test_document(markdown, "table_cell_line_breaks.docx")
        assert doc is not None
        # Find the table
        table = doc.tables[0]
        # First data row, first cell should have 2 paragraphs
        cell_0_0 = table.cell(1, 0)
        assert len(cell_0_0.paragraphs) == 2
        assert cell_0_0.paragraphs[0].text == "Line one"
        assert cell_0_0.paragraphs[1].text == "Line two"
        # First data row, second cell should have 1 paragraph
        cell_0_1 = table.cell(1, 1)
        assert len(cell_0_1.paragraphs) == 1
        # Second data row, second cell should have 3 paragraphs (A, B, C)
        cell_1_1 = table.cell(2, 1)
        assert len(cell_1_1.paragraphs) == 3
        assert cell_1_1.paragraphs[0].text == "A"
        assert cell_1_1.paragraphs[1].text == "B"
        assert cell_1_1.paragraphs[2].text == "C"


# =============================================================================
# Inline Formatting Tests
# =============================================================================

class TestInlineFormatting:
    """Tests for inline markdown formatting."""

    def test_bold_text(self):
        """Test bold text conversion."""
        markdown = "This is **bold** text."
        doc = save_test_document(markdown, "format_bold.docx")
        assert doc is not None

    def test_italic_text(self):
        """Test italic text conversion."""
        markdown = "This is *italic* text."
        doc = save_test_document(markdown, "format_italic.docx")
        assert doc is not None

    def test_inline_code(self):
        """Test inline code conversion."""
        markdown = "Use the `print()` function."
        doc = save_test_document(markdown, "format_code.docx")
        assert doc is not None

    def test_hyperlink(self):
        """Test hyperlink conversion."""
        markdown = "Visit [our website](https://example.com) for more info."
        doc = save_test_document(markdown, "format_link.docx")
        assert doc is not None

    def test_mixed_formatting(self):
        """Test multiple formatting types in one paragraph."""
        markdown = "This has **bold**, *italic*, `code`, and [link](https://test.com)."
        doc = save_test_document(markdown, "format_mixed.docx")
        assert doc is not None

    def test_nested_formatting(self):
        """Test nested formatting (bold containing italic)."""
        markdown = "This is **bold with *italic* inside**."
        doc = save_test_document(markdown, "format_nested.docx")
        assert doc is not None

    def test_escaped_characters(self):
        """Test escaped markdown characters."""
        markdown = r"This has \*asterisks\* and \**double asterisks\**."
        doc = save_test_document(markdown, "format_escaped.docx")
        assert doc is not None


# =============================================================================
# HTML Entity Tests
# =============================================================================


class TestHtmlEntities:
    """Tests for HTML entity decoding in inline text."""

    def test_nbsp_decoded(self):
        """Test that &nbsp; becomes a non-breaking space character."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("Hello&nbsp;World", p)
        assert p.text == "Hello\u00a0World"

    def test_typographic_dashes(self):
        """Test en-dash and em-dash entity decoding."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("2020&ndash;2025 &mdash; a range", p)
        assert "\u2013" in p.text  # en-dash
        assert "\u2014" in p.text  # em-dash

    def test_ellipsis(self):
        """Test &hellip; becomes … character."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("Wait for it&hellip;", p)
        assert p.text == "Wait for it\u2026"

    def test_symbols(self):
        """Test &copy; &reg; &trade; are decoded."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("&copy; 2026 Brand&trade; Product&reg;", p)
        assert "\u00a9" in p.text  # ©
        assert "\u2122" in p.text  # ™
        assert "\u00ae" in p.text  # ®

    def test_smart_quotes(self):
        """Test smart quote entities are decoded."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("&ldquo;Hello&rdquo; and &lsquo;Hi&rsquo;", p)
        assert "\u201c" in p.text  # "
        assert "\u201d" in p.text  # "
        assert "\u2018" in p.text  # '
        assert "\u2019" in p.text  # '

    def test_math_symbols(self):
        """Test &times; &divide; &plusmn; &deg; are decoded."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("5 &times; 3 &divide; 2 &plusmn; 1&deg;", p)
        assert "\u00d7" in p.text  # ×
        assert "\u00f7" in p.text  # ÷
        assert "\u00b1" in p.text  # ±
        assert "\u00b0" in p.text  # °

    def test_euro_sign(self):
        """Test &euro; is decoded."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("Price: &euro;100", p)
        assert p.text == "Price: \u20ac100"

    def test_bullet(self):
        """Test &bull; is decoded."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("Item A &bull; Item B", p)
        assert "\u2022" in p.text

    def test_dangerous_entities_not_decoded(self):
        """Test that &lt; &gt; &amp; are NOT decoded (would break markdown)."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("&lt;center&gt; &amp; &gt;quote", p)
        assert "&lt;" in p.text
        assert "&gt;" in p.text
        assert "&amp;" in p.text

    def test_entities_with_formatting(self):
        """Test entities work correctly alongside markdown formatting."""
        doc = Document()
        p = doc.add_paragraph()
        parse_inline_formatting("**Price:&nbsp;&euro;50** &mdash; *tax&nbsp;incl.*", p)
        assert "\u00a0" in p.text
        assert "\u20ac" in p.text
        assert "\u2014" in p.text


# =============================================================================
# Block Quote Tests
# =============================================================================

class TestBlockQuotes:
    """Tests for block quote conversion."""

    def test_simple_quote(self):
        """Test simple block quote."""
        markdown = "> This is a quoted text."
        doc = save_test_document(markdown, "quote_simple.docx")
        assert doc is not None

    def test_quote_with_formatting(self):
        """Test block quote with inline formatting."""
        markdown = "> This quote has **bold** and *italic* text."
        doc = save_test_document(markdown, "quote_formatted.docx")
        assert doc is not None


# =============================================================================
# Complex Document Tests
# =============================================================================

class TestComplexDocuments:
    """Tests for complex documents combining multiple elements."""

    def test_full_document(self):
        """Test a complete document with all elements."""
        markdown = """# Project Report

## Executive Summary

This report provides a **comprehensive analysis** of the project status.

## Key Findings

The following points summarize our findings:

- Revenue increased by **15%**
- Customer satisfaction improved to *92%*
- New features deployed successfully

## Data Overview

| Metric | Q1 | Q2 | Q3 |
|--------|----|----|-----|
| Sales | 100 | 120 | 150 |
| Users | 500 | 600 | 800 |

## Next Steps

1. Expand into new markets
2. Invest in R&D
3. Focus on customer retention

> "The best way to predict the future is to create it." - Peter Drucker

## Conclusion

Visit [our dashboard](https://example.com/dashboard) for live updates.
"""
        doc = save_test_document(markdown, "complex_full_document.docx")
        assert doc is not None

    def test_legal_contract_style(self):
        """Test legal contract style document with numbered sections."""
        markdown = """# SERVICE AGREEMENT

1. PARTIES
   - This agreement is between Company A and Company B.
   - Both parties agree to the following terms.

2. SERVICES
   - Company A will provide consulting services.
   - Services include analysis, recommendations, and implementation support.

3. PAYMENT TERMS
   - Payment is due within 30 days of invoice.
   - Late payments incur a 1.5% monthly fee.

4. CONFIDENTIALITY
   - Both parties agree to maintain confidentiality.
   - This obligation survives termination of the agreement.
"""
        doc = save_test_document(markdown, "complex_contract.docx")
        assert doc is not None

    def test_technical_documentation(self):
        """Test technical documentation style."""
        markdown = """# API Documentation

## Authentication

All API requests require authentication using an API key.

Use the `Authorization` header:

> Authorization: Bearer YOUR_API_KEY

## Endpoints

### GET /users

Returns a list of users.

| Parameter | Type | Required | Description |
|-----------|------|----------|-------------|
| page | integer | No | Page number |
| limit | integer | No | Items per page |

### POST /users

Creates a new user.

**Request Body:**

- `name` - User's full name
- `email` - User's email address
- `role` - User's role (*admin*, *user*, or *guest*)
"""
        doc = save_test_document(markdown, "complex_api_docs.docx")
        assert doc is not None


# =============================================================================
# Edge Cases
# =============================================================================

class TestEdgeCases:
    """Tests for edge cases and special scenarios."""

    def test_empty_content(self):
        """Test with empty content."""
        markdown = ""
        doc = save_test_document(markdown, "edge_empty_content.docx")
        assert doc is not None

    def test_only_whitespace(self):
        """Test with only whitespace."""
        markdown = "   \n\n   \n"
        doc = save_test_document(markdown, "edge_only_whitespace.docx")
        assert doc is not None

    def test_multiple_empty_lines(self):
        """Test preservation of multiple empty lines."""
        markdown = """First paragraph.


Third paragraph (after two empty lines).
"""
        doc = save_test_document(markdown, "edge_empty_lines.docx")
        assert doc is not None

    def test_unicode_content(self):
        """Test with unicode characters."""
        markdown = """# Vícejazyčný dokument

Příliš žluťoučký kůň úpěl ďábelské ódy.

日本語テキスト

Emoji: 👋 🌍 ✨
"""
        doc = save_test_document(markdown, "edge_unicode.docx")
        assert doc is not None

    def test_long_paragraph(self):
        """Test with very long paragraph."""
        long_text = "Lorem ipsum dolor sit amet. " * 50
        markdown = f"# Long Document\n\n{long_text}"
        doc = save_test_document(markdown, "edge_long_paragraph.docx")
        assert doc is not None

    def test_special_xml_characters(self):
        """Test with characters that need XML escaping."""
        markdown = "This has < and > and & characters."
        doc = save_test_document(markdown, "edge_xml_chars.docx")
        assert doc is not None

    def test_line_breaks(self):
        """Test soft line breaks (two spaces at end)."""
        markdown = """This is line one.  
This is line two (same paragraph).  
This is line three.
"""
        doc = save_test_document(markdown, "edge_line_breaks.docx")
        assert doc is not None


# =============================================================================
# Regression Tests for helpers.py changes
# =============================================================================

class TestHelpersRegression:
    """Regression tests for helpers.py functionality used by base tool."""

    def test_parse_inline_formatting_plain(self):
        """Test parse_inline_formatting with plain text."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Plain text", para)
        assert para.text == "Plain text"

    def test_parse_inline_formatting_bold(self):
        """Test parse_inline_formatting with bold."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Text with **bold** word", para)
        assert "bold" in para.text
        bold_runs = [r for r in para.runs if r.bold]
        assert len(bold_runs) > 0

    def test_parse_inline_formatting_italic(self):
        """Test parse_inline_formatting with italic."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Text with *italic* word", para)
        assert "italic" in para.text
        italic_runs = [r for r in para.runs if r.italic]
        assert len(italic_runs) > 0

    def test_parse_inline_formatting_code(self):
        """Test parse_inline_formatting with inline code."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Use `code` here", para)
        assert "code" in para.text
        code_runs = [r for r in para.runs if r.font.name == "Courier New"]
        assert len(code_runs) > 0

    def test_parse_inline_formatting_link(self):
        """Test parse_inline_formatting with hyperlink."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("Visit [link](https://example.com)", para)
        # Check that hyperlink element exists
        hyperlinks = para._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) > 0

    def test_parse_inline_formatting_nested(self):
        """Test parse_inline_formatting with nested formatting."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("This is **bold with *italic* inside**", para)

        # Should have runs with both bold and italic
        bold_italic_runs = [r for r in para.runs if r.bold and r.italic]
        assert len(bold_italic_runs) > 0

    def test_parse_inline_formatting_multiple_bold(self):
        """Test parse_inline_formatting with multiple bold sections."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("**First** and **second** bold", para)

        bold_runs = [r for r in para.runs if r.bold and r.text.strip()]
        assert len(bold_runs) >= 2


# =============================================================================
# Comprehensive Visual Test
# =============================================================================

class TestVisualInspection:
    """Comprehensive test for manual visual inspection of generated documents.

    This test creates a single document with ALL supported markdown features
    for easy visual verification in Word/LibreOffice.

    Output: tests/output/docx/VISUAL_INSPECTION_comprehensive.docx
    """

    def test_comprehensive_visual_document(self):
        """Generate a comprehensive document for visual inspection.

        This document includes ALL supported features:
        - All heading levels (H1-H6)
        - Paragraphs with various inline formatting
        - Bold, italic, bold+italic (***text***)
        - Strikethrough (~~text~~) and underline (__text__)
        - Inline code (`code`)
        - Hyperlinks ([text](url))
        - Nested formatting (bold inside italic, italic inside bold)
        - Ordered and unordered lists (including nested)
        - Tables with formatting
        - Block quotes
        - Page breaks (---) and horizontal lines (***)
        - Text alignment (<center>, <div align="right">, etc.)
        - Line breaks (two trailing spaces)
        - Escaped characters
        - Unicode and special characters
        - Images (with fallback for invalid URL)
        """
        markdown = (
            "# Comprehensive Visual Inspection Document\n"
            "\n"
            "This document is designed for **manual visual inspection** to verify that all markdown\n"
            "features are correctly converted to Word format. Open this file in Microsoft Word or\n"
            "LibreOffice Writer to check the formatting.\n"
            "\n"
            "***\n"
            "\n"
            "## 1. Heading Levels\n"
            "\n"
            "### Heading Level 3\n"
            "\n"
            "#### Heading Level 4\n"
            "\n"
            "##### Heading Level 5\n"
            "\n"
            "###### Heading Level 6\n"
            "\n"
            "***\n"
            "\n"
            "## 2. Inline Formatting\n"
            "\n"
            "This paragraph contains **bold text**, *italic text*, and ***bold italic text***.\n"
            "You can also use `inline code` for technical terms like `print()` or `variable_name`.\n"
            "\n"
            "Here is a [hyperlink to example.com](https://example.com) and another\n"
            "[link to Google](https://www.google.com).\n"
            "\n"
            "Mixed formatting: **bold with *nested italic* inside** and *italic with **nested bold** inside*.\n"
            "\n"
            "***\n"
            "\n"
            "## 3. Strikethrough and Underline\n"
            "\n"
            "This has ~~strikethrough text~~ that should appear with a line through it.\n"
            "\n"
            "This has __underlined text__ that should appear underlined.\n"
            "\n"
            "Mixed: ~~deleted~~ and __added__ in the same paragraph.\n"
            "\n"
            "Combined with bold: **~~bold strikethrough~~** and **__bold underline__**.\n"
            "\n"
            "Combined with italic: *~~italic strikethrough~~* and *__italic underline__*.\n"
            "\n"
            "***\n"
            "\n"
            "## 4. Escaped Characters\n"
            "\n"
            "These should appear as literal characters, not formatting:\n"
            "\n"
            r"\*not italic\* and \**not bold\** and \`not code\`." "\n"
            "\n"
            "***\n"
            "\n"
            "## 5. Unordered Lists\n"
            "\n"
            "Simple bullet list:\n"
            "\n"
            "- First item\n"
            "- Second item with **bold** text\n"
            "- Third item with *italic* text\n"
            "- Fourth item with `code`\n"
            "- Fifth item with [link](https://example.com)\n"
            "- Sixth item with ~~strikethrough~~ and __underline__\n"
            "\n"
            "Nested bullet list:\n"
            "\n"
            "- Main item 1\n"
            "   - Sub-item 1.1\n"
            "   - Sub-item 1.2\n"
            "      - Deep nested item\n"
            "   - Sub-item 1.3\n"
            "- Main item 2\n"
            "   - Sub-item 2.1\n"
            "\n"
            "Different markers (should all render as bullets):\n"
            "\n"
            "* Asterisk item 1\n"
            "* Asterisk item 2\n"
            "\n"
            "+ Plus item 1\n"
            "+ Plus item 2\n"
            "\n"
            "***\n"
            "\n"
            "## 6. Ordered Lists\n"
            "\n"
            "Simple numbered list:\n"
            "\n"
            "1. First step\n"
            "2. Second step with **important** info\n"
            "3. Third step with *emphasis*\n"
            "4. Fourth step with `code snippet`\n"
            "\n"
            "Nested numbered list:\n"
            "\n"
            "1. Main step 1\n"
            "   1. Sub-step 1.1\n"
            "   2. Sub-step 1.2\n"
            "2. Main step 2\n"
            "   1. Sub-step 2.1\n"
            "   2. Sub-step 2.2\n"
            "   3. Sub-step 2.3\n"
            "3. Main step 3\n"
            "\n"
            "***\n"
            "\n"
            "## 7. Mixed List Types\n"
            "\n"
            "Shopping list:\n"
            "\n"
            "- Apples\n"
            "- Bananas\n"
            "- Oranges\n"
            "\n"
            "Preparation steps:\n"
            "\n"
            "1. Wash the fruit\n"
            "2. Cut into pieces\n"
            "3. Serve and enjoy\n"
            "\n"
            "***\n"
            "\n"
            "## 8. Tables\n"
            "\n"
            "### Simple Table\n"
            "\n"
            "| Name | Age | City |\n"
            "|------|-----|------|\n"
            "| John | 25 | New York |\n"
            "| Jane | 30 | Los Angeles |\n"
            "| Bob | 35 | Chicago |\n"
            "\n"
            "### Table with Formatting\n"
            "\n"
            "| Feature | Description | Status |\n"
            "|---------|-------------|--------|\n"
            "| **Bold Feature** | This feature is *very important* | Active |\n"
            "| *Italic Feature* | Contains `code` elements | Pending |\n"
            "| ~~Removed~~ | Was __underlined__ | Archived |\n"
            "| Regular Feature | Visit [docs](https://docs.example.com) | Complete |\n"
            "\n"
            "### Table with Alignment\n"
            "\n"
            "| Left Aligned | Center Aligned | Right Aligned |\n"
            "|:-------------|:--------------:|--------------:|\n"
            "| L1 | C1 | R1 |\n"
            "| L2 | C2 | R2 |\n"
            "| L3 | C3 | R3 |\n"
            "\n"
            "***\n"
            "\n"
            "## 9. Block Quotes\n"
            "\n"
            "> This is a simple block quote.\n"
            "\n"
            "> This block quote contains **bold** and *italic* formatting.\n"
            "\n"
            "> This quote has ~~strikethrough~~ and __underline__ too.\n"
            "\n"
            '> "The best way to predict the future is to create it." - Peter Drucker\n'
            "\n"
            "***\n"
            "\n"
            "## 10. Text Alignment\n"
            "\n"
            "<center>This text should be centered.</center>\n"
            "\n"
            '<div align="right">This text should be right-aligned.</div>\n'
            "\n"
            '<div align="justify">This text should be justified. Lorem ipsum dolor sit amet, '
            "consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore "
            "magna aliqua.</div>\n"
            "\n"
            '<div align="left">This text should be left-aligned (explicit).</div>\n'
            "\n"
            "Multi-line centered block:\n"
            "\n"
            "<center>\n"
            "Company Name Inc.\n"
            "123 Main Street\n"
            "City, Country 12345\n"
            "</center>\n"
            "\n"
            "Multi-line right-aligned block:\n"
            "\n"
            '<div align="right">\n'
            "Date: 2026-02-20\n"
            "Reference: DOC-2026-001\n"
            "</div>\n"
            "\n"
            "***\n"
            "\n"
            "## 11. Unicode and Special Characters\n"
            "\n"
            "### Czech Text\n"
            "Příliš žluťoučký kůň úpěl ďábelské ódy.\n"
            "\n"
            "### German Text\n"
            "Größe, Müller, Straße, Übung\n"
            "\n"
            "### Japanese Text\n"
            "こんにちは世界 (Hello World)\n"
            "\n"
            "### Emoji\n"
            "Hello 👋 World 🌍 Stars ⭐✨ Check ✓ Heart ❤️\n"
            "\n"
            "### Special XML Characters\n"
            "5 > 3 and 2 < 4 and A & B\n"
            "\n"
            "***\n"
            "\n"
            "## 12. Line Breaks\n"
            "\n"
            "This is line one.  \n"
            "This is line two (same paragraph, soft break).  \n"
            "This is line three (still same paragraph).\n"
            "\n"
            "***\n"
            "\n"
            "## 13. Page Break\n"
            "\n"
            "The next element is a page break (---). Content after it should start on a new page.\n"
            "\n"
            "---\n"
            "\n"
            "## 14. After the Page Break\n"
            "\n"
            "This section should appear on a new page (after the --- page break above).\n"
            "\n"
            "***\n"
            "\n"
            "## 15. Images\n"
            "\n"
            "Below is an image reference (will show error placeholder since URL is invalid):\n"
            "\n"
            "![Sample Image](https://invalid-test-domain.test/sample.png)\n"
            "\n"
            "***\n"
            "\n"
            "## 16. Complex Paragraph\n"
            "\n"
            "This paragraph demonstrates **multiple formatting options** combined together.\n"
            "We have *italic text*, `inline code`, and [hyperlinks](https://example.com).\n"
            "You can even have **bold with *nested italic*** or *italic with **nested bold***.\n"
            "Also ~~strikethrough~~ and __underline__ mixed with **bold** and *italic*.\n"
            "Special characters like < > & are properly escaped.\n"
            "\n"
            "***\n"
            "\n"
            "## 17. Technical Documentation Style\n"
            "\n"
            "### API Endpoint: GET /users\n"
            "\n"
            "Returns a list of users.\n"
            "\n"
            "**Parameters:**\n"
            "\n"
            "| Parameter | Type | Required | Description |\n"
            "|-----------|------|----------|-------------|\n"
            "| `page` | integer | No | Page number (default: 1) |\n"
            "| `limit` | integer | No | Items per page (default: 20) |\n"
            "| `sort` | string | No | Sort field |\n"
            "\n"
            "**Example Response:**\n"
            "\n"
            "> The response includes user data in JSON format.\n"
            "\n"
            "***\n"
            "\n"
            "## 18. Legal Document Style\n"
            "\n"
            "1. PARTIES\n"
            "   - This agreement is between **Company A** and **Company B**.\n"
            "   - Both parties agree to the terms below.\n"
            "\n"
            "2. TERMS AND CONDITIONS\n"
            "   - All payments due within *30 days*.\n"
            "   - Late payments incur a `1.5%` monthly fee.\n"
            "\n"
            "3. CONFIDENTIALITY\n"
            "   - Both parties maintain strict confidentiality.\n"
            "   - See [Privacy Policy](https://example.com/privacy) for details.\n"
            "\n"
            "***\n"
            "\n"
            "## Conclusion\n"
            "\n"
            "This document contains **all** supported markdown elements:\n"
            "- **Bold**, *italic*, ***bold italic***\n"
            "- ~~Strikethrough~~ and __underline__\n"
            "- `Inline code` and [hyperlinks](https://example.com)\n"
            "- Headings (H1-H6), lists, tables, block quotes\n"
            "- Page breaks (---) and horizontal lines (***)\n"
            "- Text alignment (center, right, justify, left)\n"
            "- Line breaks, escaped characters, Unicode, images\n"
            "\n"
            "If you can read this and all formatting above appears correct, the markdown-to-Word\n"
            "conversion is working properly! 🎉\n"
            "\n"
            "**Document generated for visual inspection purposes.**\n"
            "\n"
            "*Last updated: February 2026*\n"
        )
        doc = save_test_document(markdown, "VISUAL_INSPECTION_comprehensive.docx")
        assert doc is not None

        # ----- Basic sanity checks -----
        assert len(doc.paragraphs) > 50, "Document should have many paragraphs"

        # ----- Text content presence -----
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Comprehensive Visual Inspection" in full_text
        assert "bold text" in full_text
        assert "italic text" in full_text
        assert "bold italic text" in full_text
        assert "strikethrough text" in full_text
        assert "underlined text" in full_text
        assert "First item" in full_text
        assert "žluťoučký" in full_text  # Czech unicode
        assert "こんにちは" in full_text  # Japanese
        assert "Größe" in full_text  # German unicode
        assert "👋" in full_text  # Emoji
        assert "5 > 3" in full_text or "5 > 3" in full_text  # Special XML characters
        assert "After the Page Break" in full_text  # Page break section
        assert "not italic" in full_text  # Escaped characters rendered as literal text

        # Multi-line alignment block content
        assert "Company Name Inc." in full_text, "Multi-line centered block content"
        assert "123 Main Street" in full_text, "Multi-line centered block content"
        assert "Date: 2026-02-20" in full_text, "Multi-line right-aligned block content"
        assert "Reference: DOC-2026-001" in full_text, "Multi-line right-aligned block content"

        # List content
        assert "Sub-item 1.1" in full_text, "Nested bullet list content"
        assert "Asterisk item 1" in full_text, "Asterisk marker list content"
        assert "Plus item 1" in full_text, "Plus marker list content"
        assert "Sub-step 1.1" in full_text, "Nested ordered list content"
        assert "Main step 3" in full_text, "Ordered list content"

        # ----- Heading levels -----
        headings = [(p.style.name, p.text) for p in doc.paragraphs
                    if p.style.name.startswith('Heading')]
        heading_levels = set(h[0] for h in headings)
        assert 'Heading 1' in heading_levels, "Should have H1 headings"
        assert 'Heading 2' in heading_levels, "Should have H2 headings"
        assert 'Heading 3' in heading_levels, "Should have H3 headings"
        assert 'Heading 4' in heading_levels, "Should have H4 headings"
        assert 'Heading 5' in heading_levels, "Should have H5 headings"
        assert 'Heading 6' in heading_levels, "Should have H6 headings"

        # ----- Inline formatting runs -----
        all_runs = [r for p in doc.paragraphs for r in p.runs]

        # Bold-only runs
        bold_only_runs = [r for r in all_runs if r.bold and not r.italic]
        assert len(bold_only_runs) > 0, "Should have bold-only runs"
        assert any("bold text" in r.text for r in bold_only_runs), \
            "Bold-only runs should contain 'bold text'"

        # Italic-only runs
        italic_only_runs = [r for r in all_runs if r.italic and not r.bold]
        assert len(italic_only_runs) > 0, "Should have italic-only runs"
        assert any("italic text" in r.text for r in italic_only_runs), \
            "Italic-only runs should contain 'italic text'"

        # Strikethrough runs
        strike_runs = [r for r in all_runs if r.font.strike]
        assert len(strike_runs) > 0, "Should have strikethrough runs"
        assert any("strikethrough" in r.text or "deleted" in r.text or "Removed" in r.text
                    for r in strike_runs), "Strikethrough should contain expected text"

        # Underline runs
        underline_runs = [r for r in all_runs if r.font.underline]
        assert len(underline_runs) > 0, "Should have underline runs"
        assert any("underlined" in r.text or "added" in r.text or "underline" in r.text.lower()
                    for r in underline_runs), "Underline should contain expected text"

        # Bold+italic runs (from ***bold italic text***)
        bold_italic_runs = [r for r in all_runs if r.bold and r.italic]
        assert len(bold_italic_runs) > 0, "Should have bold+italic runs"
        assert any("bold italic" in r.text for r in bold_italic_runs), \
            "bold+italic runs should contain 'bold italic' text"

        # Nested formatting: bold containing italic (**bold with *nested italic* inside**)
        nested_italic_in_bold = [r for r in all_runs if r.bold and r.italic
                                 and "nested italic" in r.text]
        assert len(nested_italic_in_bold) > 0, \
            "Should have bold+italic run from nested **bold with *italic* inside**"

        # Nested formatting: italic containing bold (*italic with **nested bold** inside*)
        # This should produce italic-only runs and bold+italic runs
        nested_bi_runs = [r for r in all_runs if r.bold and r.italic and "nested bold" in r.text]
        assert len(nested_bi_runs) > 0, \
            "Should have bold+italic run from nested *italic with **bold** inside*"

        # Bold + strikethrough (from **~~bold strikethrough~~**)
        bold_strike_runs = [r for r in all_runs if r.bold and r.font.strike]
        assert len(bold_strike_runs) > 0, "Should have bold+strikethrough runs"
        assert any("bold strikethrough" in r.text for r in bold_strike_runs), \
            "Bold+strikethrough run should contain 'bold strikethrough'"

        # Bold + underline (from **__bold underline__**)
        bold_underline_runs = [r for r in all_runs if r.bold and r.font.underline]
        assert len(bold_underline_runs) > 0, "Should have bold+underline runs"
        assert any("bold underline" in r.text for r in bold_underline_runs), \
            "Bold+underline run should contain 'bold underline'"

        # Italic + strikethrough (from *~~italic strikethrough~~*)
        italic_strike_runs = [r for r in all_runs if r.italic and r.font.strike]
        assert len(italic_strike_runs) > 0, "Should have italic+strikethrough runs"
        assert any("italic strikethrough" in r.text for r in italic_strike_runs), \
            "Italic+strikethrough run should contain 'italic strikethrough'"

        # Italic + underline (from *__italic underline__*)
        italic_underline_runs = [r for r in all_runs if r.italic and r.font.underline]
        assert len(italic_underline_runs) > 0, "Should have italic+underline runs"
        assert any("italic underline" in r.text for r in italic_underline_runs), \
            "Italic+underline run should contain 'italic underline'"

        # Code runs (inline code with Courier New font)
        code_runs = [r for r in all_runs if r.font.name == "Courier New"]
        assert len(code_runs) > 0, "Should have code runs"
        code_texts = [r.text for r in code_runs]
        assert any("print()" in t for t in code_texts), "Code runs should contain 'print()'"
        assert any("variable_name" in t for t in code_texts), "Code runs should contain 'variable_name'"
        assert any("inline code" == t or "code" in t.lower() for t in code_texts), \
            "Code runs should contain code-related text"

        # ----- Text alignment -----
        centered = [p for p in doc.paragraphs
                    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and p.text.strip()]
        assert len(centered) > 0, "Should have centered paragraphs"
        assert any("centered" in p.text.lower() for p in centered), \
            "Centered paragraphs should contain expected text"
        # Multi-line center block should produce multiple centered paragraphs
        # (inline centered + Company Name + 123 Main Street + City)
        assert len(centered) >= 4, \
            "Should have >=4 centered paragraphs (inline + multi-line block)"

        right_aligned = [p for p in doc.paragraphs
                         if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT and p.text.strip()]
        assert len(right_aligned) > 0, "Should have right-aligned paragraphs"
        assert any("right-aligned" in p.text.lower() for p in right_aligned), \
            "Right-aligned paragraphs should contain expected text"
        # Multi-line right block: inline + Date + Reference = >=3
        assert len(right_aligned) >= 3, \
            "Should have >=3 right-aligned paragraphs (inline + multi-line block)"

        justified = [p for p in doc.paragraphs
                     if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and p.text.strip()]
        assert len(justified) > 0, "Should have justified paragraphs"
        assert any("justified" in p.text.lower() or "lorem ipsum" in p.text.lower()
                    for p in justified), "Justified paragraph should contain expected text"

        left_aligned = [p for p in doc.paragraphs
                        if p.alignment == WD_ALIGN_PARAGRAPH.LEFT and p.text.strip()]
        assert len(left_aligned) > 0, "Should have explicit left-aligned paragraphs"
        assert any("left-aligned" in p.text.lower() for p in left_aligned), \
            "Left-aligned paragraph should contain expected text"

        # ----- Block quotes -----
        quote_paragraphs = [p for p in doc.paragraphs if p.style.name == 'Quote']
        assert len(quote_paragraphs) >= 4, \
            "Should have at least 4 block quote paragraphs"
        quote_texts = [p.text for p in quote_paragraphs]
        assert any("simple block quote" in t for t in quote_texts), \
            "Block quotes should include simple quote"
        assert any("bold" in t and "italic" in t for t in quote_texts), \
            "Block quotes should include formatted quote"
        assert any("Peter Drucker" in t for t in quote_texts), \
            "Block quotes should include attribution quote"

        # Block quote with formatting runs
        quote_runs = [r for p in doc.paragraphs if p.style.name == 'Quote'
                      for r in p.runs]
        quote_bold = [r for r in quote_runs if r.bold]
        assert len(quote_bold) > 0, "Block quotes should have bold runs"
        quote_italic = [r for r in quote_runs if r.italic]
        assert len(quote_italic) > 0, "Block quotes should have italic runs"
        quote_strike = [r for r in quote_runs if r.font.strike]
        assert len(quote_strike) > 0, "Block quotes should have strikethrough runs"
        quote_underline = [r for r in quote_runs if r.font.underline]
        assert len(quote_underline) > 0, "Block quotes should have underline runs"

        # ----- Lists (verify styles) -----
        bullet_paras = [p for p in doc.paragraphs
                        if p.style.name.startswith('List Bullet')]
        assert len(bullet_paras) >= 6, \
            f"Should have >=6 bullet list paragraphs, got {len(bullet_paras)}"
        # Nested bullets should use List Bullet 2 or 3
        nested_bullets = [p for p in bullet_paras if p.style.name != 'List Bullet']
        assert len(nested_bullets) > 0, "Should have nested bullet list paragraphs"

        number_paras = [p for p in doc.paragraphs
                        if p.style.name.startswith('List Number')]
        assert len(number_paras) >= 4, \
            f"Should have >=4 numbered list paragraphs, got {len(number_paras)}"
        # Nested numbers should use List Number 2 or 3
        nested_numbers = [p for p in number_paras if p.style.name != 'List Number']
        assert len(nested_numbers) > 0, "Should have nested numbered list paragraphs"

        # List items with formatting
        list_runs = [r for p in doc.paragraphs
                     if p.style.name.startswith('List Bullet') or
                     p.style.name.startswith('List Number')
                     for r in p.runs]
        list_bold = [r for r in list_runs if r.bold]
        assert len(list_bold) > 0, "List items should have bold runs"
        list_italic = [r for r in list_runs if r.italic]
        assert len(list_italic) > 0, "List items should have italic runs"
        list_code = [r for r in list_runs if r.font.name == "Courier New"]
        assert len(list_code) > 0, "List items should have code runs"
        list_strike = [r for r in list_runs if r.font.strike]
        assert len(list_strike) > 0, "List items should have strikethrough runs"
        list_underline = [r for r in list_runs if r.font.underline]
        assert len(list_underline) > 0, "List items should have underline runs"

        # ----- Tables -----
        assert len(doc.tables) >= 3, "Document should have at least 3 tables"
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
        assert "John" in table_text, "Simple table data"
        assert "Jane" in table_text, "Simple table data"
        assert "Bob" in table_text, "Simple table data"
        assert "New York" in table_text, "Simple table city data"
        assert "Bold Feature" in table_text, "Table with formatting"
        assert "Italic Feature" in table_text, "Table with formatting"
        assert "Removed" in table_text, "Table with strikethrough"
        assert "Left Aligned" in table_text or "L1" in table_text, "Alignment table"

        # Table cells should have inline formatting (bold, italic, code, strikethrough, underline)
        table_runs = [r for table in doc.tables
                      for row in table.rows for cell in row.cells
                      for p in cell.paragraphs for r in p.runs]
        table_bold = [r for r in table_runs if r.bold]
        assert len(table_bold) > 0, "Table cells should have bold formatting"
        table_italic = [r for r in table_runs if r.italic]
        assert len(table_italic) > 0, "Table cells should have italic formatting"
        table_code = [r for r in table_runs if r.font.name == "Courier New"]
        assert len(table_code) > 0, "Table cells should have code formatting"
        table_strike = [r for r in table_runs if r.font.strike]
        assert len(table_strike) > 0, "Table cells should have strikethrough formatting"
        table_underline = [r for r in table_runs if r.font.underline]
        assert len(table_underline) > 0, "Table cells should have underline formatting"

        # Table hyperlinks
        table_xml = "".join(
            table._tbl.xml for table in doc.tables
        )
        assert 'w:hyperlink' in table_xml, "Table should have hyperlinks"

        # ----- Page break (---) in XML -----
        xml = doc.element.xml
        assert 'w:br' in xml or 'type="page"' in xml, "Should have page break"

        # ----- Horizontal line (***) as w:pBdr -----
        assert 'w:pBdr' in xml, "Should have horizontal line borders"
        # Multiple horizontal lines (***) used as section separators
        pBdr_count = xml.count('w:pBdr')
        assert pBdr_count >= 2, f"Should have multiple horizontal lines, got {pBdr_count}"

        # ----- Image error placeholder (invalid URL) -----
        assert "Image could not be loaded" in full_text, "Should have image error placeholder"
        assert "invalid-test-domain.test" in full_text, \
            "Image error should include the URL"

        # ----- Line breaks (two trailing spaces -> w:br) -----
        # The "Line Breaks" section uses trailing spaces to produce soft breaks
        line_break_count = xml.count('<w:br/>')
        assert line_break_count >= 2, \
            f"Should have >=2 soft line breaks from trailing double-spaces, got {line_break_count}"

        # ----- Hyperlinks -----
        hyperlink_count = xml.count('w:hyperlink')
        assert hyperlink_count >= 4, \
            f"Should have at least 4 hyperlinks (example, google, docs, privacy), got {hyperlink_count}"

        # ----- Escaped characters (should NOT have formatting) -----
        # The escaped line should render as literal text with *, **, `
        escaped_para = None
        for p in doc.paragraphs:
            if "not italic" in p.text and "not bold" in p.text:
                escaped_para = p
                break
        assert escaped_para is not None, "Should find the escaped characters paragraph"
        # None of the runs in the escaped paragraph should have formatting
        for r in escaped_para.runs:
            if "not italic" in r.text:
                assert not r.italic, "Escaped *text* should NOT be italic"
            if "not bold" in r.text:
                assert not r.bold, "Escaped **text** should NOT be bold"

    def test_comprehensive_visual_with_metadata_and_toc(self):
        """Generate a comprehensive document with metadata, TOC, header and footer.

        Output: tests/output/docx/VISUAL_INSPECTION_metadata_toc.docx
        """
        markdown = """# Chapter 1: Introduction

This is the introduction chapter. It demonstrates that the **Table of Contents**,
document **metadata**, and **headers/footers** with page numbers work correctly.

## 1.1 Background

Some background information with *italic* and **bold** formatting.

## 1.2 Objectives

1. Verify TOC generation
2. Verify metadata fields
3. Verify header and footer with page numbers

---

# Chapter 2: Features

## 2.1 Strikethrough and Underline

~~Old feature~~ replaced by __new feature__.

## 2.2 Text Alignment

<center>Centered heading text</center>

<div align="right">Right-aligned date: 2026-02-20</div>

## 2.3 Bold Italic

This is ***bold and italic*** text together.

---

# Chapter 3: Conclusion

All features verified. Check the header, footer (with page numbers), TOC,
and document properties (title, author, subject) in Word.

**End of document.**
"""
        doc = create_word_document(
            markdown,
            title="Visual Inspection Document",
            author="Test Suite",
            subject="Comprehensive Feature Verification",
            header_text="Visual Inspection Report",
            footer_text="Page {page} of {pages}",
            include_toc=True,
        )
        output_path = OUTPUT_DIR / "VISUAL_INSPECTION_metadata_toc.docx"
        doc.save(str(output_path))
        print(f"Saved: {output_path}")

        # Verify metadata
        assert doc.core_properties.title == "Visual Inspection Document"
        assert doc.core_properties.author == "Test Suite"
        assert doc.core_properties.subject == "Comprehensive Feature Verification"

        # Verify TOC field exists
        xml = doc.element.xml
        assert 'TOC' in xml, "Should have TOC field"
        assert 'updateFields' in doc.settings.element.xml, "Should have updateFields setting"

        # Verify header
        header = doc.sections[0].header
        header_text = "\n".join([p.text for p in header.paragraphs])
        assert "Visual Inspection Report" in header_text

        # Verify footer with page fields
        footer = doc.sections[0].footer
        footer_xml = footer._element.xml
        assert 'PAGE' in footer_xml, "Footer should contain PAGE field"
        assert 'NUMPAGES' in footer_xml, "Footer should contain NUMPAGES field"

        # Verify content
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Table of Contents" in full_text
        assert "Chapter 1" in full_text
        assert "Chapter 2" in full_text
        assert "Chapter 3" in full_text


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])


# =============================================================================
# Page Break Tests (Feature 1)
# =============================================================================

class TestPageBreaks:
    """Tests for page break (---) conversion."""

    def test_page_break_creates_break_element(self):
        """Test that --- creates a page break (w:br type=page)."""
        markdown = "First page content\n\n---\n\nSecond page content"
        doc = save_test_document(markdown, "page_break_basic.docx")
        xml = doc.element.xml
        assert 'w:br' in xml or 'lastRenderedPageBreak' in xml or 'type="page"' in xml

    def test_multiple_page_breaks(self):
        """Test multiple page breaks in a document."""
        markdown = "Page 1\n\n---\n\nPage 2\n\n---\n\nPage 3"
        doc = save_test_document(markdown, "page_break_multiple.docx")
        assert doc is not None

    def test_page_break_with_long_dashes(self):
        """Test that ---- (more than 3 dashes) also creates page break."""
        markdown = "Before\n\n----\n\nAfter"
        doc = save_test_document(markdown, "page_break_long_dashes.docx")
        assert doc is not None


# =============================================================================
# Horizontal Line Tests (Feature 1)
# =============================================================================

class TestHorizontalLines:
    """Tests for horizontal line (***) conversion."""

    def test_horizontal_line_creates_border(self):
        """Test that *** creates a horizontal line with bottom border."""
        markdown = "Text above\n\n***\n\nText below"
        doc = save_test_document(markdown, "hline_basic.docx")
        xml = doc.element.xml
        assert 'w:pBdr' in xml
        assert 'w:bottom' in xml

    def test_multiple_horizontal_lines(self):
        """Test multiple horizontal lines."""
        markdown = "Section 1\n\n***\n\nSection 2\n\n***\n\nSection 3"
        doc = save_test_document(markdown, "hline_multiple.docx")
        assert doc is not None


# =============================================================================
# Image Tests (Feature 2)
# =============================================================================

class TestImages:
    """Tests for image (![alt](url)) conversion."""

    def test_image_with_invalid_url(self):
        """Test image with invalid URL creates error placeholder."""
        markdown = "![Test](https://invalid-domain-that-does-not-exist.test/img.png)"
        doc = save_test_document(markdown, "image_invalid_url.docx")
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Image could not be loaded" in full_text

    def test_image_placeholder_text(self):
        """Test that failed image includes the URL in error text."""
        url = "https://nonexistent.test/photo.jpg"
        markdown = f"![Photo]({url})"
        doc = save_test_document(markdown, "image_placeholder.docx")
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert url in full_text


# =============================================================================
# Text Alignment Tests (Feature 3)
# =============================================================================

class TestTextAlignment:
    """Tests for text alignment via HTML tags."""

    def test_center_alignment(self):
        """Test <center>text</center> creates centered paragraph."""
        markdown = "<center>Centered text</center>"
        doc = save_test_document(markdown, "align_center.docx")
        # Find the paragraph with the centered text
        for p in doc.paragraphs:
            if "Centered text" in p.text:
                assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
                break
        else:
            pytest.fail("Centered text paragraph not found")

    def test_right_alignment(self):
        """Test <div align="right">text</div> creates right-aligned paragraph."""
        markdown = '<div align="right">Right aligned text</div>'
        doc = save_test_document(markdown, "align_right.docx")
        for p in doc.paragraphs:
            if "Right aligned" in p.text:
                assert p.alignment == WD_ALIGN_PARAGRAPH.RIGHT
                break
        else:
            pytest.fail("Right-aligned text paragraph not found")

    def test_justify_alignment(self):
        """Test <div align="justify">text</div> creates justified paragraph."""
        markdown = '<div align="justify">Justified text content</div>'
        doc = save_test_document(markdown, "align_justify.docx")
        for p in doc.paragraphs:
            if "Justified text" in p.text:
                assert p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
                break
        else:
            pytest.fail("Justified text paragraph not found")

    def test_multiline_center_block(self):
        """Test multi-line <center> block."""
        markdown = "<center>\nCompany Name\nStreet Address\nCity, Country\n</center>"
        doc = save_test_document(markdown, "align_multiline_center.docx")
        centered_paragraphs = [p for p in doc.paragraphs
                               if p.alignment == WD_ALIGN_PARAGRAPH.CENTER
                               and p.text.strip()]
        assert len(centered_paragraphs) >= 3

    def test_multiline_div_right_block(self):
        """Test multi-line <div align="right"> block."""
        markdown = '<div align="right">\nDate: 2026-02-20\nRef: ABC-123\n</div>'
        doc = save_test_document(markdown, "align_multiline_right.docx")
        right_paragraphs = [p for p in doc.paragraphs
                            if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT
                            and p.text.strip()]
        assert len(right_paragraphs) >= 2


# =============================================================================
# Document Metadata Tests (Feature 5)
# =============================================================================

class TestDocumentMetadata:
    """Tests for document metadata (title, author, subject)."""

    def test_metadata_title(self):
        """Test that document title is set."""
        doc = create_word_document("# Test", title="My Document Title")
        assert doc.core_properties.title == "My Document Title"

    def test_metadata_author(self):
        """Test that document author is set."""
        doc = create_word_document("# Test", author="John Doe")
        assert doc.core_properties.author == "John Doe"

    def test_metadata_subject(self):
        """Test that document subject is set."""
        doc = create_word_document("# Test", subject="Annual Report 2026")
        assert doc.core_properties.subject == "Annual Report 2026"

    def test_metadata_all_fields(self):
        """Test that all metadata fields are set together."""
        doc = create_word_document(
            "# Report",
            title="Annual Report",
            author="Jane Smith",
            subject="Financial Overview"
        )
        assert doc.core_properties.title == "Annual Report"
        assert doc.core_properties.author == "Jane Smith"
        assert doc.core_properties.subject == "Financial Overview"

    def test_metadata_none_not_set(self):
        """Test that None metadata does not overwrite defaults."""
        doc = create_word_document("# Test")
        # Should not throw, core_properties should exist
        assert doc.core_properties is not None


# =============================================================================
# Header/Footer Tests (Feature 6)
# =============================================================================

class TestHeadersFooters:
    """Tests for document headers and footers with page numbers."""

    def test_header_plain_text(self):
        """Test simple header text."""
        doc = create_word_document("# Test", header_text="Company Report")
        header = doc.sections[0].header
        header_text = "\n".join([p.text for p in header.paragraphs])
        assert "Company Report" in header_text

    def test_footer_plain_text(self):
        """Test simple footer text."""
        doc = create_word_document("# Test", footer_text="Confidential")
        footer = doc.sections[0].footer
        footer_text = "\n".join([p.text for p in footer.paragraphs])
        assert "Confidential" in footer_text

    def test_footer_with_page_number(self):
        """Test footer with {page} token inserts PAGE field."""
        doc = create_word_document("# Test", footer_text="Page {page} of {pages}")
        footer = doc.sections[0].footer
        xml = footer._element.xml
        assert 'PAGE' in xml
        assert 'NUMPAGES' in xml

    def test_header_and_footer_together(self):
        """Test both header and footer set simultaneously."""
        doc = create_word_document(
            "# Test",
            header_text="Header Text",
            footer_text="Footer Text"
        )
        header_text = "\n".join([p.text for p in doc.sections[0].header.paragraphs])
        footer_text = "\n".join([p.text for p in doc.sections[0].footer.paragraphs])
        assert "Header Text" in header_text
        assert "Footer Text" in footer_text

    def test_header_footer_saved(self):
        """Test header and footer are preserved when saving."""
        doc = save_test_document("# Report\n\nSome content.", "header_footer_saved.docx")
        # Just verify doc is created successfully - manual inspection for formatting
        assert doc is not None


# =============================================================================
# Table of Contents Tests (Feature 7)
# =============================================================================

class TestTableOfContents:
    """Tests for Table of Contents insertion."""

    def test_toc_field_exists(self):
        """Test that TOC field elements are present in document XML."""
        doc = create_word_document(
            "# Chapter 1\n\nContent\n\n## Section 1.1\n\nMore content",
            include_toc=True
        )
        xml = doc.element.xml
        assert 'TOC' in xml
        assert 'fldChar' in xml or 'fldSimple' in xml

    def test_toc_heading_exists(self):
        """Test that 'Table of Contents' heading is added."""
        doc = create_word_document("# Test Heading\n\nContent", include_toc=True)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Table of Contents" in full_text

    def test_toc_update_fields_setting(self):
        """Test that updateFields setting is added to document."""
        doc = create_word_document("# Test", include_toc=True)
        xml = doc.settings.element.xml
        assert 'updateFields' in xml

    def test_toc_saved_document(self):
        """Test that TOC document saves correctly."""
        doc = save_test_document(
            "# Chapter 1\n\nIntro\n\n## Section 1.1\n\nDetails\n\n# Chapter 2\n\nConclusion",
            "toc_document.docx"
        )
        assert doc is not None


# =============================================================================
# Underline and Strikethrough Tests (Feature 8)
# =============================================================================

class TestUnderlineStrikethrough:
    """Tests for ~~strikethrough~~ and __underline__ formatting."""

    def test_strikethrough(self):
        """Test ~~text~~ creates strikethrough run."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("This is ~~deleted~~ text", para)
        strike_runs = [r for r in para.runs if r.font.strike]
        assert len(strike_runs) > 0
        assert any("deleted" in r.text for r in strike_runs)

    def test_underline(self):
        """Test __text__ creates underlined run."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("This is __important__ text", para)
        underline_runs = [r for r in para.runs if r.font.underline]
        assert len(underline_runs) > 0
        assert any("important" in r.text for r in underline_runs)

    def test_bold_and_underline(self):
        """Test **__text__** creates bold and underlined run."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("**__bold and underlined__**", para)
        bu_runs = [r for r in para.runs if r.bold and r.font.underline]
        assert len(bu_runs) > 0

    def test_mixed_strikethrough_underline(self):
        """Test mixed ~~old~~ __new__ text."""
        doc = Document()
        para = doc.add_paragraph()
        parse_inline_formatting("This is ~~old~~ __new__ text", para)
        strike_runs = [r for r in para.runs if r.font.strike]
        underline_runs = [r for r in para.runs if r.font.underline]
        assert len(strike_runs) > 0
        assert len(underline_runs) > 0

    def test_strikethrough_in_document(self):
        """Test strikethrough in full document creation."""
        markdown = "This has ~~deleted text~~ in it."
        doc = save_test_document(markdown, "format_strikethrough.docx")
        assert doc is not None

    def test_underline_in_document(self):
        """Test underline in full document creation."""
        markdown = "This has __underlined text__ in it."
        doc = save_test_document(markdown, "format_underline.docx")
        assert doc is not None

    def test_superscript(self):
        """Test superscript formatting with ^text^."""
        markdown = "E = mc^2^ is famous."
        doc = save_test_document(markdown, "format_superscript.docx")
        assert doc is not None
        para = doc.paragraphs[-1]
        runs = para.runs
        # Find the superscript run
        super_run = next(r for r in runs if r.text == '2')
        assert super_run.font.superscript is True

    def test_subscript(self):
        """Test subscript formatting with ~text~."""
        markdown = "Water is H~2~O."
        doc = save_test_document(markdown, "format_subscript.docx")
        assert doc is not None
        para = doc.paragraphs[-1]
        runs = para.runs
        sub_run = next(r for r in runs if r.text == '2')
        assert sub_run.font.subscript is True

    def test_highlight(self):
        """Test highlight formatting with ==text==."""
        from docx.enum.text import WD_COLOR_INDEX
        markdown = "This is ==very important== text."
        doc = save_test_document(markdown, "format_highlight.docx")
        assert doc is not None
        para = doc.paragraphs[-1]
        runs = para.runs
        hl_run = next(r for r in runs if r.text == 'very important')
        assert hl_run.font.highlight_color == WD_COLOR_INDEX.YELLOW

    def test_superscript_subscript_combined(self):
        """Test super and subscript in the same paragraph."""
        markdown = "x^2^ + y~i~ = z"
        doc = save_test_document(markdown, "format_super_sub_combined.docx")
        assert doc is not None
        para = doc.paragraphs[-1]
        runs = para.runs
        super_run = next(r for r in runs if r.text == '2')
        sub_run = next(r for r in runs if r.text == 'i')
        assert super_run.font.superscript is True
        assert sub_run.font.subscript is True


