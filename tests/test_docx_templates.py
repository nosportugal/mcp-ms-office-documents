"""Tests for dynamic DOCX template creation.

These tests create actual .docx files and save them to disk for manual inspection.
Output files are saved to tests/output/docx/ directory.

Test coverage:
- Basic placeholder replacement
- Markdown formatting (bold, italic, code, links)
- Multiple placeholders in one paragraph
- Placeholders in tables
- Placeholders in headers/footers
- Empty/missing placeholders
- Special characters and unicode
- Template registration from YAML
"""

import sys
from pathlib import Path

# Add project root to path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

import pytest
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_tools.dynamic_docx_tools import (
    _replace_placeholders_in_paragraph,
    _replace_placeholders_in_document,
    find_docx_template_by_name,
)
from docx_tools.helpers import contains_block_markdown

# Output directory for test files
OUTPUT_DIR = Path(__file__).parent / "output" / "docx"
TEMPLATES_DIR = Path(__file__).parent / "templates"


@pytest.fixture(scope="module", autouse=True)
def setup_output_dir():
    """Create output directory if it doesn't exist."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    yield


def save_document(doc: Document, filename: str) -> Path:
    """Save document to output directory and return path."""
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
    return output_path


def create_test_document_with_placeholder(placeholder: str, font_size: int = 11) -> Document:
    """Create a simple test document with a single placeholder."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run(placeholder)
    run.font.size = Pt(font_size)
    return doc


def get_paragraph_runs_info(paragraph) -> list:
    """Get info about all runs in a paragraph including formatting."""
    runs_info = []
    for run in paragraph.runs:
        info = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "font_name": run.font.name,
        }
        runs_info.append(info)
    return runs_info


# =============================================================================
# Basic Placeholder Replacement Tests
# =============================================================================

class TestBasicPlaceholderReplacement:
    """Tests for basic placeholder replacement functionality."""

    def test_simple_placeholder_replacement(self):
        """Test replacing a simple placeholder with plain text."""
        doc = create_test_document_with_placeholder("Hello {{name}}!")
        context = {"name": "World"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "basic_01_simple_replacement.docx")
        assert path.exists()

        # Verify content
        doc2 = Document(path)
        assert "World" in doc2.paragraphs[0].text
        assert "{{name}}" not in doc2.paragraphs[0].text

    def test_multiple_placeholders_same_paragraph(self):
        """Test replacing multiple placeholders in the same paragraph."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Dear {{title}} {{name}}, welcome to {{company}}!")

        context = {
            "title": "Mr.",
            "name": "Smith",
            "company": "Acme Corp"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "basic_02_multiple_placeholders.docx")
        assert path.exists()

        doc2 = Document(path)
        text = doc2.paragraphs[0].text
        assert "Mr." in text
        assert "Smith" in text
        assert "Acme Corp" in text

    def test_placeholder_at_start(self):
        """Test placeholder at the start of paragraph."""
        doc = create_test_document_with_placeholder("{{greeting}} everyone!")
        context = {"greeting": "Hello"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "basic_03_placeholder_at_start.docx")
        assert path.exists()

    def test_placeholder_at_end(self):
        """Test placeholder at the end of paragraph."""
        doc = create_test_document_with_placeholder("Best regards, {{signature}}")
        context = {"signature": "John Doe"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "basic_04_placeholder_at_end.docx")
        assert path.exists()

    def test_placeholder_only(self):
        """Test paragraph containing only a placeholder."""
        doc = create_test_document_with_placeholder("{{content}}")
        context = {"content": "This is the entire content."}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "basic_05_placeholder_only.docx")
        assert path.exists()

    def test_empty_replacement_value(self):
        """Test replacing placeholder with empty string."""
        doc = create_test_document_with_placeholder("Hello {{optional}}World")
        context = {"optional": ""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "basic_06_empty_replacement.docx")
        assert path.exists()

    def test_missing_placeholder_in_context(self):
        """Test that missing placeholders are left unchanged."""
        doc = create_test_document_with_placeholder("Hello {{unknown_placeholder}}!")
        context = {"other": "value"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "basic_07_missing_placeholder.docx")
        assert path.exists()

        # Verify placeholder is still there
        doc2 = Document(path)
        assert "{{unknown_placeholder}}" in doc2.paragraphs[0].text


# =============================================================================
# Markdown Formatting Tests
# =============================================================================

class TestMarkdownFormatting:
    """Tests for markdown formatting support in placeholder values."""

    def test_bold_text(self):
        """Test bold text formatting (**text**)."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "This is **bold** text"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_01_bold.docx")
        assert path.exists()

        # Verify formatting
        doc2 = Document(path)
        runs = get_paragraph_runs_info(doc2.paragraphs[0])
        bold_runs = [r for r in runs if r["bold"] and r["text"].strip()]
        assert len(bold_runs) > 0
        assert any("bold" in r["text"] for r in bold_runs)

    def test_italic_text(self):
        """Test italic text formatting (*text*)."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "This is *italic* text"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_02_italic.docx")
        assert path.exists()

        # Verify formatting
        doc2 = Document(path)
        runs = get_paragraph_runs_info(doc2.paragraphs[0])
        italic_runs = [r for r in runs if r["italic"] and r["text"].strip()]
        assert len(italic_runs) > 0

    def test_inline_code(self):
        """Test inline code formatting (`code`)."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "Use the `print()` function"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_03_code.docx")
        assert path.exists()

        # Verify code uses monospace font
        doc2 = Document(path)
        runs = get_paragraph_runs_info(doc2.paragraphs[0])
        code_runs = [r for r in runs if r["font_name"] == "Courier New"]
        assert len(code_runs) > 0

    def test_hyperlink(self):
        """Test hyperlink formatting [text](url)."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "Visit [our website](https://example.com) for more info"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_04_hyperlink.docx")
        assert path.exists()

        # Verify hyperlink exists - check the XML for hyperlink element
        doc2 = Document(path)
        para = doc2.paragraphs[0]
        hyperlinks = para._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
        assert len(hyperlinks) > 0

    def test_mixed_formatting(self):
        """Test multiple formatting types in one value."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {
            "message": "This has **bold**, *italic*, `code`, and [link](https://test.com)"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_05_mixed_formatting.docx")
        assert path.exists()

    def test_nested_formatting(self):
        """Test nested formatting - bold containing italic."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "This is **bold with *italic* inside**"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_06_nested.docx")
        assert path.exists()

        # Verify nested formatting is applied correctly
        doc2 = Document(path)
        runs = get_paragraph_runs_info(doc2.paragraphs[0])

        # Should have: plain "This is ", bold "bold with ", bold+italic "italic", bold " inside"
        bold_runs = [r for r in runs if r["bold"] and r["text"].strip()]
        italic_runs = [r for r in runs if r["italic"] and r["text"].strip()]
        bold_italic_runs = [r for r in runs if r["bold"] and r["italic"] and r["text"].strip()]

        assert len(bold_runs) >= 1, "Should have bold text"
        assert len(italic_runs) >= 1, "Should have italic text"
        assert len(bold_italic_runs) >= 1, "Should have bold+italic text (nested)"
        assert any("italic" in r["text"] for r in bold_italic_runs), "The word 'italic' should be bold+italic"

    def test_nested_italic_containing_bold(self):
        """Test nested formatting - italic containing bold."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "This is *italic with **bold** inside*"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_06b_nested_italic_bold.docx")
        assert path.exists()

        # Verify nested formatting
        doc2 = Document(path)
        runs = get_paragraph_runs_info(doc2.paragraphs[0])
        bold_italic_runs = [r for r in runs if r["bold"] and r["italic"] and r["text"].strip()]
        assert len(bold_italic_runs) >= 1, "Should have bold+italic text (nested)"

    def test_multiple_bold_sections(self):
        """Test multiple bold sections in one value."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "**First** and **second** and **third** bold words"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_07_multiple_bold.docx")
        assert path.exists()

    def test_formatting_preserves_base_font(self):
        """Test that markdown formatting preserves the original font size."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("{{message}}")
        run.font.size = Pt(14)  # Set specific font size

        context = {"message": "Text with **bold** formatting"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "markdown_08_preserve_font.docx")
        assert path.exists()


# =============================================================================
# Table Placeholder Tests
# =============================================================================

class TestTablePlaceholders:
    """Tests for placeholders in tables."""

    def test_placeholder_in_table_cell(self):
        """Test placeholder replacement in a table cell."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'

        # Add headers
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"

        # Add placeholders
        table.cell(1, 0).text = "{{field_name}}"
        table.cell(1, 1).text = "{{field_value}}"

        context = {
            "field_name": "Company",
            "field_value": "Acme Corp"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "table_01_simple.docx")
        assert path.exists()

    def test_markdown_in_table_cell(self):
        """Test markdown formatting in table cells."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'

        table.cell(0, 0).text = "Feature"
        table.cell(0, 1).text = "Description"
        table.cell(1, 0).text = "{{feature}}"
        table.cell(1, 1).text = "{{description}}"

        context = {
            "feature": "**Bold Feature**",
            "description": "This is *very* important"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "table_02_markdown.docx")
        assert path.exists()

    def test_multiple_placeholders_in_table(self):
        """Test multiple placeholders across table cells."""
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'

        # Fill with placeholders
        for i in range(3):
            for j in range(3):
                table.cell(i, j).text = f"{{{{cell_{i}_{j}}}}}"

        context = {f"cell_{i}_{j}": f"R{i}C{j}" for i in range(3) for j in range(3)}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "table_03_multiple.docx")
        assert path.exists()


# =============================================================================
# Header and Footer Tests
# =============================================================================

class TestHeaderFooterPlaceholders:
    """Tests for placeholders in headers and footers."""

    def test_placeholder_in_header(self):
        """Test placeholder replacement in document header."""
        doc = Document()

        # Add a section with header
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.add_run("Document: {{doc_title}}")

        # Add body content
        doc.add_paragraph("Body content")

        context = {"doc_title": "Annual Report 2026"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "header_01_simple.docx")
        assert path.exists()

    def test_placeholder_in_footer(self):
        """Test placeholder replacement in document footer."""
        doc = Document()

        # Add a section with footer
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.add_run("© {{year}} {{company}}")

        # Add body content
        doc.add_paragraph("Body content")

        context = {
            "year": "2026",
            "company": "Test Company"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "footer_01_simple.docx")
        assert path.exists()

    def test_markdown_in_header(self):
        """Test markdown formatting in header."""
        doc = Document()

        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.add_run("{{header_content}}")

        doc.add_paragraph("Body content")

        context = {"header_content": "**Important Document** - *Confidential*"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "header_02_markdown.docx")
        assert path.exists()


# =============================================================================
# Unicode and Special Characters Tests
# =============================================================================

class TestUnicodeAndSpecialCharacters:
    """Tests for unicode and special character handling."""

    def test_unicode_replacement_value(self):
        """Test replacement with unicode characters."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "Příliš žluťoučký kůň úpěl ďábelské ódy"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "unicode_01_czech.docx")
        assert path.exists()

        doc2 = Document(path)
        assert "žluťoučký" in doc2.paragraphs[0].text

    def test_emoji_in_replacement(self):
        """Test replacement with emoji characters."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "Hello 👋 World 🌍!"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "unicode_02_emoji.docx")
        assert path.exists()

    def test_special_xml_characters(self):
        """Test replacement with characters that need XML escaping."""
        doc = create_test_document_with_placeholder("{{message}}")
        context = {"message": "5 > 3 and 2 < 4 and A & B"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "unicode_03_xml_special.docx")
        assert path.exists()

    def test_multiline_replacement(self):
        """Test replacement with newline characters."""
        doc = create_test_document_with_placeholder("{{address}}")
        context = {"address": "123 Main Street\nApartment 4B\nNew York, NY 10001"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "unicode_04_multiline.docx")
        assert path.exists()

    def test_japanese_characters(self):
        """Test replacement with Japanese characters."""
        doc = create_test_document_with_placeholder("{{greeting}}")
        context = {"greeting": "こんにちは世界"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "unicode_05_japanese.docx")
        assert path.exists()


# =============================================================================
# Complex Document Tests
# =============================================================================

class TestComplexDocuments:
    """Tests for complex document scenarios."""

    def test_formal_letter_template(self):
        """Test a complete formal letter template."""
        doc = Document()

        # Date (right-aligned)
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run("{{date}}")

        doc.add_paragraph()  # Empty line

        # Recipient
        doc.add_paragraph().add_run("{{recipient_name}}").bold = True
        doc.add_paragraph().add_run("{{recipient_address}}")

        doc.add_paragraph()

        # Subject
        subject_para = doc.add_paragraph()
        run = subject_para.add_run("Subject: ")
        run.bold = True
        subject_para.add_run("{{subject}}")

        doc.add_paragraph()

        # Salutation
        doc.add_paragraph().add_run("{{salutation}}")

        doc.add_paragraph()

        # Body
        doc.add_paragraph().add_run("{{body}}")

        doc.add_paragraph()

        # Closing
        doc.add_paragraph().add_run("{{closing}}")

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature
        doc.add_paragraph().add_run("{{sender_name}}").bold = True
        doc.add_paragraph().add_run("{{sender_title}}").italic = True

        context = {
            "date": "January 4, 2026",
            "recipient_name": "Jan Novák",
            "recipient_address": "Hlavní 123\n110 00 Praha 1",
            "subject": "Partnership Proposal",
            "salutation": "Dear Mr. Novák,",
            "body": "I am writing to propose a **strategic partnership** between our companies. "
                    "This opportunity would allow us to *leverage synergies* and create "
                    "significant value for both parties.\n\n"
                    "Please visit [our website](https://example.com) for more details.",
            "closing": "Best regards,",
            "sender_name": "John Smith",
            "sender_title": "Chief Executive Officer"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "complex_01_formal_letter.docx")
        assert path.exists()

    def test_invoice_template(self):
        """Test an invoice-like template with tables."""
        doc = Document()

        # Header
        header_para = doc.add_paragraph()
        header_para.add_run("INVOICE").bold = True

        doc.add_paragraph()

        # Invoice info table
        info_table = doc.add_table(rows=2, cols=2)
        info_table.cell(0, 0).text = "Invoice Number:"
        info_table.cell(0, 1).text = "{{invoice_number}}"
        info_table.cell(1, 0).text = "Date:"
        info_table.cell(1, 1).text = "{{invoice_date}}"

        doc.add_paragraph()

        # Bill to section
        doc.add_paragraph().add_run("Bill To:").bold = True
        doc.add_paragraph().add_run("{{client_name}}")
        doc.add_paragraph().add_run("{{client_address}}")

        doc.add_paragraph()

        # Items table
        items_table = doc.add_table(rows=3, cols=3)
        items_table.style = 'Table Grid'

        # Headers
        items_table.cell(0, 0).text = "Description"
        items_table.cell(0, 1).text = "Quantity"
        items_table.cell(0, 2).text = "Amount"

        # Items
        items_table.cell(1, 0).text = "{{item1_desc}}"
        items_table.cell(1, 1).text = "{{item1_qty}}"
        items_table.cell(1, 2).text = "{{item1_amount}}"

        items_table.cell(2, 0).text = "{{item2_desc}}"
        items_table.cell(2, 1).text = "{{item2_qty}}"
        items_table.cell(2, 2).text = "{{item2_amount}}"

        context = {
            "invoice_number": "INV-2026-0001",
            "invoice_date": "January 4, 2026",
            "client_name": "Acme Corporation",
            "client_address": "456 Business Ave\nSuite 100\nNew York, NY",
            "item1_desc": "**Consulting Services**",
            "item1_qty": "40 hours",
            "item1_amount": "$4,000.00",
            "item2_desc": "*Software License*",
            "item2_qty": "1",
            "item2_amount": "$500.00",
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "complex_02_invoice.docx")
        assert path.exists()

    def test_report_with_sections(self):
        """Test a report-style document with multiple sections."""
        doc = Document()

        # Title
        title = doc.add_heading("{{report_title}}", level=0)

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph().add_run("{{executive_summary}}")

        # Key Findings
        doc.add_heading("Key Findings", level=1)
        doc.add_paragraph().add_run("{{key_findings}}")

        # Recommendations
        doc.add_heading("Recommendations", level=1)
        doc.add_paragraph().add_run("{{recommendations}}")

        # Conclusion
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph().add_run("{{conclusion}}")

        context = {
            "report_title": "Q4 2025 Analysis Report",
            "executive_summary": "This report provides a **comprehensive analysis** of Q4 2025 performance. "
                                 "Key metrics show *significant improvement* across all sectors.",
            "key_findings": "1. Revenue increased by **15%**\n"
                           "2. Customer satisfaction improved to *92%*\n"
                           "3. New market entry was `successful`",
            "recommendations": "Based on our analysis, we recommend:\n"
                              "- Expand into [new markets](https://example.com/markets)\n"
                              "- Invest in **R&D**\n"
                              "- Focus on *customer retention*",
            "conclusion": "The quarter exceeded expectations. Continued focus on innovation "
                         "will drive future growth."
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "complex_03_report.docx")
        assert path.exists()


# =============================================================================
# Edge Cases and Error Handling Tests
# =============================================================================

class TestEdgeCases:
    """Tests for edge cases and error handling."""

    def test_placeholder_with_underscore(self):
        """Test placeholder names with underscores."""
        doc = create_test_document_with_placeholder("{{first_name}} {{last_name}}")
        context = {
            "first_name": "John",
            "last_name": "Doe"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "edge_01_underscore_names.docx")
        assert path.exists()

    def test_placeholder_with_numbers(self):
        """Test placeholder names with numbers."""
        doc = create_test_document_with_placeholder("{{item1}} {{item2}} {{item3}}")
        context = {
            "item1": "First",
            "item2": "Second",
            "item3": "Third"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "edge_02_numbered_names.docx")
        assert path.exists()

    def test_very_long_replacement(self):
        """Test replacement with very long text."""
        doc = create_test_document_with_placeholder("{{content}}")
        long_text = "Lorem ipsum dolor sit amet. " * 100
        context = {"content": long_text}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "edge_03_long_text.docx")
        assert path.exists()

    def test_consecutive_placeholders(self):
        """Test placeholders directly next to each other."""
        doc = create_test_document_with_placeholder("{{first}}{{second}}{{third}}")
        context = {
            "first": "A",
            "second": "B",
            "third": "C"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "edge_04_consecutive.docx")
        assert path.exists()

    def test_empty_document_with_placeholder(self):
        """Test document with only a placeholder."""
        doc = create_test_document_with_placeholder("{{only_content}}")
        context = {"only_content": "This is the entire document content."}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "edge_05_empty_doc.docx")
        assert path.exists()

    def test_placeholder_not_in_runs(self):
        """Test behavior when paragraph has no runs (edge case)."""
        doc = Document()
        para = doc.add_paragraph()
        # Paragraph exists but has no runs - should not crash

        context = {"test": "value"}
        _replace_placeholders_in_paragraph(para, context)

        path = save_document(doc, "edge_06_no_runs.docx")
        assert path.exists()

    def test_triple_brace_placeholder(self):
        """Test triple-brace mustache syntax {{{name}}}."""
        doc = create_test_document_with_placeholder("{{{raw_content}}}")
        context = {"raw_content": "Content with **formatting**"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "edge_07_triple_brace.docx")
        assert path.exists()


# =============================================================================
# YAML Registration Tests
# =============================================================================

class TestYAMLRegistration:
    """Tests for YAML-based template registration."""

    @pytest.fixture
    def sample_yaml_config(self, tmp_path):
        """Create a sample YAML config file for testing."""
        yaml_content = """
templates:
  - name: test_letter
    description: Test letter template
    docx_path: letter_template.docx
    annotations:
      title: Test Letter
    args:
      - name: recipient
        type: string
        description: Recipient name
        required: true
      - name: message
        type: string
        description: Message content
        required: true
      - name: optional_note
        type: string
        description: Optional note
        required: false
        default: ""
"""
        yaml_path = tmp_path / "test_config.yaml"
        yaml_path.write_text(yaml_content, encoding="utf-8")
        return yaml_path

    def test_yaml_loading(self, sample_yaml_config):
        """Test that YAML config is loaded correctly."""
        import yaml
        content = sample_yaml_config.read_text(encoding="utf-8")
        config = yaml.safe_load(content)

        assert "templates" in config
        assert len(config["templates"]) == 1
        assert config["templates"][0]["name"] == "test_letter"

    def test_find_template_in_custom_dir(self):
        """Test finding templates in custom_templates directory."""
        # This test depends on the actual letter_template.docx existing
        from docx_tools.dynamic_docx_tools import find_docx_template_by_name

        result = find_docx_template_by_name("letter_template.docx")
        # Should find the template we created earlier
        if result:
            assert "letter_template.docx" in result
            assert Path(result).exists()


# =============================================================================
# List (Bullet Points and Numbered Lists) Tests
# =============================================================================

class TestListsInPlaceholders:
    """Tests for bullet points and numbered lists in placeholder values."""

    def test_value_contains_block_content_detection_unordered(self):
        """Test detection of unordered list in value."""
        value = """Some intro text
- First item
- Second item
- Third item"""
        assert contains_block_markdown(value) is True

    def test_value_contains_block_content_detection_ordered(self):
        """Test detection of ordered list in value."""
        value = """Some intro text
1. First step
2. Second step
3. Third step"""
        assert contains_block_markdown(value) is True

    def test_value_contains_block_content_plain_text(self):
        """Test that plain text is not detected as block content."""
        value = "This is just plain text without any lists."
        assert contains_block_markdown(value) is False

    def test_value_contains_block_content_inline_formatting(self):
        """Test that inline formatting is not detected as block content."""
        value = "This has **bold** and *italic* but no lists."
        assert contains_block_markdown(value) is False

    def test_value_contains_block_content_heading(self):
        """Test detection of heading in value."""
        value = """# Main Heading
Some text here
## Sub heading"""
        assert contains_block_markdown(value) is True

    def test_simple_unordered_list(self):
        """Test placeholder with simple unordered list."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{items}}")

        context = {"items": """- Apple
- Banana
- Orange"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_01_simple_unordered.docx")
        assert path.exists()

        # Verify list items were created as separate paragraphs
        doc2 = Document(path)
        # Should have list paragraphs
        list_paragraphs = [p for p in doc2.paragraphs if p.text.strip() and
                          ('Apple' in p.text or 'Banana' in p.text or 'Orange' in p.text)]
        assert len(list_paragraphs) >= 3

    def test_simple_ordered_list(self):
        """Test placeholder with simple ordered list."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{steps}}")

        context = {"steps": """1. First step
2. Second step
3. Third step"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_02_simple_ordered.docx")
        assert path.exists()

        # Verify list items were created
        doc2 = Document(path)
        list_paragraphs = [p for p in doc2.paragraphs if p.text.strip() and
                          ('First step' in p.text or 'Second step' in p.text or 'Third step' in p.text)]
        assert len(list_paragraphs) >= 3

    def test_unordered_list_with_formatting(self):
        """Test unordered list items with markdown formatting."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{items}}")

        context = {"items": """- **Bold item**
- *Italic item*
- Item with `code`
- Item with [link](https://example.com)"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_03_unordered_formatted.docx")
        assert path.exists()

    def test_ordered_list_with_formatting(self):
        """Test ordered list items with markdown formatting."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{steps}}")

        context = {"steps": """1. **Important first step**
2. Do *something* here
3. Use `function()` to complete"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_04_ordered_formatted.docx")
        assert path.exists()

    def test_list_with_preceding_text(self):
        """Test list with text before it."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """Here are the key points:
- First point
- Second point
- Third point"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_05_with_preceding_text.docx")
        assert path.exists()

    def test_list_with_following_text(self):
        """Test list with text after it."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """- First item
- Second item
- Third item

That's all for now."""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_06_with_following_text.docx")
        assert path.exists()

    def test_mixed_list_types(self):
        """Test document with both ordered and unordered lists."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """Shopping list:
- Apples
- Bananas
- Oranges

Steps to follow:
1. Go to store
2. Buy items
3. Return home"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_07_mixed_types.docx")
        assert path.exists()

    def test_nested_unordered_list(self):
        """Test nested unordered list items."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{items}}")

        context = {"items": """- Main item 1
   - Sub item 1.1
   - Sub item 1.2
- Main item 2
   - Sub item 2.1"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_08_nested_unordered.docx")
        assert path.exists()

    def test_nested_ordered_list(self):
        """Test nested ordered list items."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{steps}}")

        context = {"steps": """1. First main step
   1. Sub-step 1.1
   2. Sub-step 1.2
2. Second main step
   1. Sub-step 2.1"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_09_nested_ordered.docx")
        assert path.exists()

    def test_list_placeholder_in_context(self):
        """Test list placeholder when there is text before and after the placeholder."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Before: {{list}} After the list.")

        context = {"list": """- Item A
- Item B
- Item C"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_10_with_context.docx")
        assert path.exists()

        # Verify "Before:" text is present
        doc2 = Document(path)
        full_text = " ".join(p.text for p in doc2.paragraphs)
        assert "Before:" in full_text
        assert "After the list." in full_text

    def test_asterisk_list_marker(self):
        """Test unordered list with asterisk marker (*)."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{items}}")

        context = {"items": """* Apple
* Banana
* Orange"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_11_asterisk_marker.docx")
        assert path.exists()

    def test_plus_list_marker(self):
        """Test unordered list with plus marker (+)."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{items}}")

        context = {"items": """+ Apple
+ Banana
+ Orange"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_12_plus_marker.docx")
        assert path.exists()

    def test_list_in_table_cell_fallback(self):
        """Test that lists in table cells fallback to inline text.

        Lists are not supported in table cells, so the value should be
        inserted as plain formatted text.
        """
        doc = Document()
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Items"
        table.cell(1, 0).text = "{{items}}"

        context = {"items": """- Item 1
- Item 2
- Item 3"""}

        # This should work without error, lists just won't be formatted as lists
        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_13_table_fallback.docx")
        assert path.exists()

    def test_complex_document_with_lists(self):
        """Test a complex document with multiple lists and formatting."""
        doc = Document()

        # Add heading
        doc.add_heading("Project Overview", level=1)

        # Add paragraph with list placeholder
        para = doc.add_paragraph()
        para.add_run("Key features: {{features}}")

        # Add another paragraph
        doc.add_paragraph().add_run("Implementation steps: {{steps}}")

        context = {
            "features": """
- **Performance** - Optimized for speed
- **Security** - Enterprise-grade protection
- **Scalability** - Grows with your needs""",
            "steps": """
1. Install the package
2. Configure settings
3. Run the setup wizard
4. Deploy to production"""
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_14_complex_document.docx")
        assert path.exists()

    def test_list_with_empty_lines(self):
        """Test list with empty lines between items."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{items}}")

        context = {"items": """- First item

- Second item

- Third item"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "list_15_with_empty_lines.docx")
        assert path.exists()


# =============================================================================
# Heading Tests in Custom Templates
# =============================================================================

class TestHeadingsInPlaceholders:
    """Tests for markdown headings in placeholder values."""

    def test_simple_heading(self):
        """Test placeholder with a simple heading."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """# Main Title
This is some body text."""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "heading_01_simple.docx")
        assert path.exists()

        # Verify heading was created
        doc2 = Document(str(path))
        # Check that heading style was applied
        heading_paragraphs = [p for p in doc2.paragraphs if p.style.name.startswith('Heading')]
        assert len(heading_paragraphs) >= 1

    def test_multiple_heading_levels(self):
        """Test placeholder with multiple heading levels."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """# Heading 1
Introduction text.

## Heading 2
More details here.

### Heading 3
Even more specific."""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "heading_02_multiple_levels.docx")
        assert path.exists()

    def test_heading_with_formatting(self):
        """Test heading with inline markdown formatting."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """# **Bold** and *italic* heading
Regular paragraph text."""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "heading_03_with_formatting.docx")
        assert path.exists()

    def test_heading_with_lists(self):
        """Test heading followed by lists."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """# Shopping List
- Apples
- Bananas
- Oranges

## Steps
1. Go to store
2. Buy items
3. Return home"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "heading_04_with_lists.docx")
        assert path.exists()

    def test_h1_to_h6_headings(self):
        """Test all heading levels from H1 to H6."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {"content": """# Heading 1
## Heading 2
### Heading 3
#### Heading 4
##### Heading 5
###### Heading 6"""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "heading_05_all_levels.docx")
        assert path.exists()

    def test_heading_in_complex_document(self):
        """Test headings in a complex document structure."""
        doc = Document()
        doc.add_heading("Document Title", level=0)
        para = doc.add_paragraph()
        para.add_run("{{sections}}")
        doc.add_paragraph("Footer text")

        context = {"sections": """# Introduction
This document covers important topics.

## Background
Some background information.

## Main Points
- Point one
- Point two
- Point three

## Conclusion
Final thoughts here."""}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "heading_06_complex_document.docx")
        assert path.exists()


# =============================================================================
# Performance Tests
# =============================================================================

class TestPerformance:
    """Performance-related tests."""

    def test_many_placeholders(self):
        """Test document with many placeholders."""
        doc = Document()

        # Create 50 paragraphs with placeholders
        for i in range(50):
            doc.add_paragraph().add_run(f"Item {{{{item_{i}}}}}: {{{{value_{i}}}}}")

        context = {}
        for i in range(50):
            context[f"item_{i}"] = f"Item {i}"
            context[f"value_{i}"] = f"Value **{i}**"

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "perf_01_many_placeholders.docx")
        assert path.exists()

    def test_large_table_with_placeholders(self):
        """Test large table with placeholders in each cell."""
        doc = Document()

        rows, cols = 10, 5
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        context = {}
        for i in range(rows):
            for j in range(cols):
                placeholder = f"{{{{cell_{i}_{j}}}}}"
                table.cell(i, j).text = placeholder
                context[f"cell_{i}_{j}"] = f"R{i}C{j}"

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "perf_02_large_table.docx")
        assert path.exists()


# =============================================================================
# Integration Tests
# =============================================================================

class TestIntegration:
    """Integration tests using actual template files."""

    def test_full_letter_workflow(self):
        """Test complete workflow of loading template and generating document."""
        from docx_tools.dynamic_docx_tools import find_docx_template_by_name

        # Find the letter template
        template_path = find_docx_template_by_name("letter_template.docx")
        if not template_path:
            pytest.skip("letter_template.docx not found in template directories")

        # Load the template
        doc = Document(template_path)

        # Define context
        context = {
            "date": "4. ledna 2026",
            "recipient_name": "Ing. Pavel Novotný",
            "recipient_address": "Technická 2\n166 27 Praha 6",
            "subject": "Nabídka spolupráce",
            "salutation": "Vážený pane inženýre,",
            "body": "dovolujeme si Vám nabídnout **exkluzivní spolupráci** v oblasti "
                    "*softwarového vývoje*. Naše společnost disponuje týmem zkušených "
                    "vývojářů a můžeme Vám pomoci s realizací Vašich projektů.\n\n"
                    "Více informací naleznete na [našich stránkách](https://example.com).",
            "closing": "S úctou,",
            "sender_name": "Mgr. Jana Svobodová",
            "sender_title": "Obchodní ředitelka"
        }

        # Replace placeholders
        _replace_placeholders_in_document(doc, context)

        # Save the result
        path = save_document(doc, "integration_01_letter.docx")
        assert path.exists()

        # Verify content
        doc2 = Document(path)
        full_text = "\n".join([p.text for p in doc2.paragraphs])
        assert "Pavel Novotný" in full_text
        assert "spolupráce" in full_text


# =============================================================================
# Markdown Tables in Placeholders Tests
# =============================================================================

class TestTablesInPlaceholders:
    """Tests for markdown table syntax in placeholder values."""

    def test_value_contains_block_content_detection_table(self):
        """Test that contains_block_markdown detects table lines."""
        table_md = "| Col1 | Col2 |\n|------|------|\n| A    | B    |"
        assert contains_block_markdown(table_md) is True

    def test_simple_table_in_placeholder(self):
        """Test replacing a placeholder with a markdown table."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{data}}")

        context = {
            "data": "| Name | Age |\n|------|-----|\n| Alice | 30 |\n| Bob | 25 |"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "table_md_01_simple.docx")
        assert path.exists()

        doc2 = Document(path)
        # Should have at least one table
        assert len(doc2.tables) >= 1
        table = doc2.tables[0]
        assert len(table.rows) == 3  # header row + 2 data rows (separator is filtered)
        assert "Name" in table.cell(0, 0).text
        assert "Alice" in table.cell(1, 0).text
        assert "Bob" in table.cell(2, 0).text

    def test_table_with_preceding_text(self):
        """Test placeholder with text before the table."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Before {{content}} After")

        context = {
            "content": "Here is the data:\n\n| H1 | H2 |\n|----|----|\n| V1 | V2 |"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "table_md_02_with_text.docx")
        assert path.exists()

        doc2 = Document(path)
        assert len(doc2.tables) >= 1

    def test_table_with_inline_formatting_in_cells(self):
        """Test markdown table with formatted cell content."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{report}}")

        context = {
            "report": "| Feature | Status |\n|---------|--------|\n| **Auth** | *Done* |\n| `API` | Pending |"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "table_md_03_formatted_cells.docx")
        assert path.exists()

        doc2 = Document(path)
        assert len(doc2.tables) >= 1

    def test_table_mixed_with_lists(self):
        """Test placeholder with both table and list content."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{content}}")

        context = {
            "content": "Summary:\n\n- Item 1\n- Item 2\n\n| Col A | Col B |\n|-------|-------|\n| X | Y |"
        }

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "table_md_04_mixed.docx")
        assert path.exists()

        doc2 = Document(path)
        assert len(doc2.tables) >= 1


# =============================================================================
# Header/Footer with Existing Template Tests
# =============================================================================

class TestHeaderFooterWithTemplate:
    """Tests for header/footer handling when the template already has headers/footers."""

    def test_set_header_preserves_alignment(self):
        """Test that set_header_footer preserves existing paragraph alignment."""
        from docx_tools.helpers import set_header_footer

        doc = Document()
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.add_run("Old header text")

        doc.add_paragraph("Body content for visual inspection.")

        set_header_footer(doc, "New header text", 'header')

        # Alignment should be preserved from template
        assert section.header.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert "New header text" in section.header.paragraphs[0].text

        path = save_document(doc, "header_tpl_01_preserves_alignment.docx")
        assert path.exists()

    def test_set_footer_all_sections(self):
        """Test that set_header_footer updates all sections."""
        from docx_tools.helpers import set_header_footer

        doc = Document()
        doc.add_paragraph("Section 1 content.")
        # Add a second section
        doc.add_section()
        doc.add_paragraph("Section 2 content.")
        assert len(doc.sections) == 2

        set_header_footer(doc, "Page {page} of {pages}", 'footer')

        # Both sections should have the footer
        for section in doc.sections:
            footer_text = section.footer.paragraphs[0].text
            # The text between fields is preserved; PAGE/NUMPAGES are fields
            assert "Page" in footer_text or len(section.footer.paragraphs[0].runs) > 0

        path = save_document(doc, "header_tpl_02_footer_all_sections.docx")
        assert path.exists()

    def test_set_header_replaces_existing_content(self):
        """Test that existing header content is replaced, not appended."""
        from docx_tools.helpers import set_header_footer

        doc = Document()
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].add_run("Original header content")

        doc.add_paragraph("Body content for visual inspection.")

        set_header_footer(doc, "Replacement header", 'header')

        header_text = section.header.paragraphs[0].text
        assert "Original header content" not in header_text
        assert "Replacement header" in header_text

        path = save_document(doc, "header_tpl_03_replaces_existing.docx")
        assert path.exists()

    def test_set_header_with_page_fields(self):
        """Test header with page number tokens."""
        from docx_tools.helpers import set_header_footer

        doc = Document()
        doc.add_paragraph("Body content for visual inspection.")

        set_header_footer(doc, "Page {page} of {pages}", 'header')

        # Should have runs with text and field elements
        p = doc.sections[0].header.paragraphs[0]
        assert len(p.runs) > 0

        path = save_document(doc, "header_tpl_04_page_fields.docx")
        assert path.exists()

    def test_placeholder_in_first_page_header(self):
        """Test placeholder replacement in first-page header of a template."""
        doc = Document()
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        # Add placeholder to first-page header
        first_header = section.first_page_header
        first_header.paragraphs[0].add_run("{{company}} - First Page")

        # Add placeholder to default header
        default_header = section.header
        default_header.paragraphs[0].add_run("{{company}} - Other Pages")

        doc.add_paragraph("Body content")

        context = {"company": "Acme Corp"}
        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "header_03_first_page.docx")
        assert path.exists()

        doc2 = Document(path)
        s = doc2.sections[0]
        assert "Acme Corp" in s.first_page_header.paragraphs[0].text
        assert "{{company}}" not in s.first_page_header.paragraphs[0].text
        assert "Acme Corp" in s.header.paragraphs[0].text
        assert "{{company}}" not in s.header.paragraphs[0].text


# =============================================================================
# Font Color Preservation Tests
# =============================================================================

class TestFontColorPreservation:
    """Tests for preserving font color from template placeholders."""

    def test_color_preserved_on_replacement(self):
        """Test that font color from the placeholder run is preserved."""
        from docx.shared import RGBColor

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("{{colored_text}}")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red

        context = {"colored_text": "This should be red"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "color_01_preserved.docx")
        assert path.exists()

        doc2 = Document(path)
        para = doc2.paragraphs[0]
        # Check that at least one run has the red color
        red_runs = [r for r in para.runs if r.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)]
        assert len(red_runs) > 0, "Replacement text should preserve red color"

    def test_color_preserved_with_inline_formatting(self):
        """Test that color is preserved even with markdown formatting."""
        from docx.shared import RGBColor

        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("{{styled}}")
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green
        run.font.size = Pt(14)

        context = {"styled": "This is **bold** and *italic*"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "color_02_with_formatting.docx")
        assert path.exists()

        doc2 = Document(path)
        para = doc2.paragraphs[0]
        green_runs = [r for r in para.runs if r.font.color.rgb == RGBColor(0x00, 0x80, 0x00)]
        assert len(green_runs) > 0, "Replacement text should preserve green color"

    def test_no_color_no_override(self):
        """Test that when no color is set, replacement runs don't get unexpected color."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{plain}}")

        context = {"plain": "No color specified"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "color_03_no_color.docx")
        assert path.exists()

        doc2 = Document(path)
        para = doc2.paragraphs[0]
        for run in para.runs:
            # Should not have a color set (None or default)
            assert run.font.color.rgb is None, "Run should not have color when template has no color"

    def test_multiple_placeholders_different_colors(self):
        """Test multiple placeholders with different colors in the same paragraph."""
        from docx.shared import RGBColor

        doc = Document()
        para = doc.add_paragraph()
        run1 = para.add_run("{{red_text}} and {{blue_text}}")
        # Note: in a real template each placeholder would be in its own run
        # For testing we create separate runs
        para.clear()
        r1 = para.add_run("{{red_text}}")
        r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        r2 = para.add_run(" and ")
        r3 = para.add_run("{{blue_text}}")
        r3.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

        context = {"red_text": "Red content", "blue_text": "Blue content"}

        _replace_placeholders_in_document(doc, context)

        path = save_document(doc, "color_04_multiple_colors.docx")
        assert path.exists()

        doc2 = Document(path)
        full_text = doc2.paragraphs[0].text
        assert "Red content" in full_text
        assert "Blue content" in full_text


# =============================================================================
# Comprehensive Visual Inspection Test
# =============================================================================

class TestVisualInspection:
    """Comprehensive test for manual visual inspection of template-based documents.

    This test creates a document with ALL supported placeholder and markdown features
    for easy visual verification in Word/LibreOffice.

    Output: tests/output/docx/VISUAL_INSPECTION_templates.docx
    """

    def test_comprehensive_template_visual_document(self):
        """Generate a comprehensive template-based document for visual inspection.

        This document demonstrates:
        - Placeholder replacement in various positions
        - All inline markdown formatting (bold, italic, code, links)
        - Block-level content (headings, lists)
        - Nested formatting
        - Tables with placeholders
        - Headers and footers
        - Unicode and special characters
        """
        doc = Document()

        # === HEADER ===
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.add_run("{{company_name}} - {{document_type}}")

        # === FOOTER ===
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.add_run("© {{year}} {{company_name}} | {{confidentiality}}")

        # === DOCUMENT TITLE ===
        title = doc.add_heading("{{document_title}}", level=0)

        # === INTRODUCTION ===
        doc.add_heading("1. Introduction", level=1)
        intro_para = doc.add_paragraph()
        intro_para.add_run("{{introduction}}")

        # === INLINE FORMATTING SECTION ===
        doc.add_heading("2. Inline Formatting Demo", level=1)
        formatting_para = doc.add_paragraph()
        formatting_para.add_run("{{formatting_demo}}")

        # === LISTS SECTION ===
        doc.add_heading("3. Lists Demo", level=1)

        doc.add_heading("3.1 Unordered List", level=2)
        unordered_para = doc.add_paragraph()
        unordered_para.add_run("{{unordered_list}}")

        doc.add_heading("3.2 Ordered List", level=2)
        ordered_para = doc.add_paragraph()
        ordered_para.add_run("{{ordered_list}}")

        doc.add_heading("3.3 Mixed Lists", level=2)
        mixed_para = doc.add_paragraph()
        mixed_para.add_run("{{mixed_lists}}")

        # === HEADINGS IN PLACEHOLDER ===
        doc.add_heading("4. Dynamic Sections", level=1)
        sections_para = doc.add_paragraph()
        sections_para.add_run("{{dynamic_sections}}")

        # === TABLE SECTION ===
        doc.add_heading("5. Data Table", level=1)
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # Headers
        table.cell(0, 0).text = "{{col1_header}}"
        table.cell(0, 1).text = "{{col2_header}}"
        table.cell(0, 2).text = "{{col3_header}}"

        # Data rows
        for i in range(1, 4):
            for j in range(3):
                table.cell(i, j).text = f"{{{{row{i}_col{j+1}}}}}"

        doc.add_paragraph()  # Spacing

        # === UNICODE SECTION ===
        doc.add_heading("6. Unicode & Special Characters", level=1)
        unicode_para = doc.add_paragraph()
        unicode_para.add_run("{{unicode_content}}")

        # === CONCLUSION ===
        doc.add_heading("7. Conclusion", level=1)
        conclusion_para = doc.add_paragraph()
        conclusion_para.add_run("{{conclusion}}")

        # === SIGNATURE ===
        doc.add_paragraph()
        doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.add_run("{{signature_block}}")

        # === CONTEXT VALUES ===
        context = {
            # Header/Footer
            "company_name": "**Acme Corporation**",
            "document_type": "Visual Inspection Report",
            "year": "2026",
            "confidentiality": "*Confidential*",

            # Title
            "document_title": "Comprehensive Template Visual Inspection",

            # Introduction
            "introduction": """This document demonstrates **all supported features** of the template placeholder system. 
It includes *inline formatting*, `code elements`, and [hyperlinks](https://example.com).

The purpose is to allow **manual visual inspection** to verify correct rendering in Word.""",

            # Formatting demo
            "formatting_demo": """Here you can see various formatting options:

- **Bold text** for emphasis
- *Italic text* for subtle emphasis
- `Inline code` for technical terms
- [Hyperlinks](https://example.com) for references
- **Bold with *nested italic* inside**
- *Italic with **nested bold** inside*

All these should render correctly in the Word document.""",

            # Unordered list
            "unordered_list": """- First bullet item
- Second item with **bold**
- Third item with *italic*
- Fourth item with `code`
- Fifth item with [link](https://example.com)
   - Nested sub-item 1
   - Nested sub-item 2""",

            # Ordered list
            "ordered_list": """1. First numbered step
2. Second step with **important** note
3. Third step with *emphasis*
4. Fourth step with `command`
   1. Sub-step 4.1
   2. Sub-step 4.2
5. Fifth and final step""",

            # Mixed lists
            "mixed_lists": """Features to implement:
- User authentication
- Data validation
- Error handling

Implementation order:
1. Design the architecture
2. Write unit tests
3. Implement features
4. Deploy to production""",

            # Dynamic sections with headings
            "dynamic_sections": """## Section A: Overview
This section provides an overview of the topic.

### Subsection A.1
Details about the first aspect.

### Subsection A.2
Details about the second aspect.

## Section B: Implementation
Here we discuss implementation details.

- Key point 1
- Key point 2
- Key point 3""",

            # Table data
            "col1_header": "**Name**",
            "col2_header": "**Role**",
            "col3_header": "**Status**",
            "row1_col1": "Jan Novák",
            "row1_col2": "*Developer*",
            "row1_col3": "Active",
            "row2_col1": "Marie Svobodová",
            "row2_col2": "*Designer*",
            "row2_col3": "Active",
            "row3_col1": "Petr Černý",
            "row3_col2": "*Manager*",
            "row3_col3": "On Leave",

            # Unicode content
            "unicode_content": """This section tests unicode character handling:

**Czech:** Příliš žluťoučký kůň úpěl ďábelské ódy.

**German:** Größe, Müller, Straße, Übung

**Japanese:** こんにちは世界 (Hello World)

**Emoji:** 👋 🌍 ⭐ ✨ ✓ ❤️ 🎉

**Special XML:** 5 > 3 and 2 < 4 and A & B""",

            # Conclusion
            "conclusion": """This document has demonstrated **all major features** of the template system:

1. Basic placeholder replacement
2. Inline markdown formatting
3. Block-level content (headings and lists)
4. Tables with placeholders
5. Headers and footers
6. Unicode and special character handling

If all elements above render correctly, the template system is working as expected! 🎉""",

            # Signature
            "signature_block": """**Prepared by:**
*Quality Assurance Team*
Acme Corporation

[Contact us](https://example.com/contact)"""
        }

        # Replace all placeholders
        _replace_placeholders_in_document(doc, context)

        # Save the document
        path = save_document(doc, "VISUAL_INSPECTION_templates.docx")
        assert path.exists()

        # Verify content
        doc2 = Document(path)
        full_text = "\n".join([p.text for p in doc2.paragraphs])

        # Basic content checks
        assert "Comprehensive Template Visual Inspection" in full_text
        assert "Bold text" in full_text
        assert "First bullet item" in full_text
        assert "žluťoučký" in full_text
        assert "こんにちは" in full_text


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])

